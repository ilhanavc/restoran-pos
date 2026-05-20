# -*- coding: utf-8 -*-
"""
Restoran POS v1.1.0 — Yazılım Gereksinimleri Belirtimi (SRS) generator
IEEE 830-1998 tabanlı SEC308 şablonuna göre üretilir.

Çalıştır: python docs/generate-srs.py
Çıktı: docs/SRS-Restoran-POS-v1.1.0.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_table(doc, header, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for r in t.rows:
            for c, w in zip(r.cells, col_widths):
                c.width = w
    return t


def page_break(doc):
    doc.add_page_break()


# ---------- Build document ----------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================================================
# KAPAK SAYFASI
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ATATÜRK ÜNİVERSİTESİ")
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("İktisadi ve İdari Bilimler Fakültesi")
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Yönetim Bilişim Sistemleri Bölümü")
run.font.size = Pt(14)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RESTORAN POS v1.1.0")
run.bold = True
run.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Yazılım Gereksinimleri Belirtimi (SRS)")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Restoran Yönetimi için Çok Modüllü Masaüstü POS Sistemi")
run.italic = True
run.font.size = Pt(12)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hazırlayan")
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("İlhan AVCI")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("220357008")
run.font.size = Pt(12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sürüm 1.1.0 — Mayıs 2026")
run.italic = True
run.font.size = Pt(11)

page_break(doc)

# ============================================================
# REVİZYON GEÇMİŞİ
# ============================================================
add_heading(doc, "Revizyon Geçmişi", level=1)
add_table(
    doc,
    ["Tarih", "Sürüm", "Açıklama", "Yazar"],
    [
        ["2026-04-01", "0.1", "İlk taslak — kapsam ve ana modüller belirlendi", "İlhan Avcı"],
        ["2026-04-15", "0.5", "Fonksiyonel gereksinimler genişletildi (12 sprint çıktısı)", "İlhan Avcı"],
        ["2026-04-30", "0.9", "İşlevsel olmayan gereksinimler + güvenlik bölümü (FAZ 0)", "İlhan Avcı"],
        ["2026-05-15", "1.1.0", "Bitirme teslim sürümü — final revizyon", "İlhan Avcı"],
    ],
    col_widths=[Cm(3), Cm(2), Cm(8.5), Cm(3.5)],
)
page_break(doc)

# ============================================================
# İÇİNDEKİLER (manuel başlık — Word'ün otomatik TOC alanı sonradan eklenir)
# ============================================================
add_heading(doc, "İçindekiler", level=1)
toc_items = [
    ("1. Giriş", "1"),
    ("    1.1 Amaç", "1"),
    ("    1.2 Kapsam", "1"),
    ("    1.3 Tanımlar ve Kısaltmalar", "2"),
    ("    1.4 Referanslar", "3"),
    ("    1.5 Genel Bakış", "3"),
    ("2. Genel Açıklama", "4"),
    ("    2.1 Ürün Perspektifi", "4"),
    ("    2.2 Ürün İşlevleri", "5"),
    ("    2.3 Kullanıcı Özellikleri", "6"),
    ("    2.4 Genel Sınırlamalar", "7"),
    ("    2.5 Varsayımlar ve Bağımlılıkları", "7"),
    ("3. Özel Gereksinimler", "8"),
    ("    3.1 Dış Arabirim Gereksinimleri", "8"),
    ("    3.2 Fonksiyonel Gereksinimler", "10"),
    ("    3.3 Kullanım Durumları", "16"),
    ("    3.4 Sınıflar / Nesneler", "18"),
    ("    3.5 İşlevsel Olmayan Gereksinimler", "19"),
    ("    3.6 Ters Gereksinimler", "21"),
    ("    3.7 Tasarım Kısıtlamaları", "21"),
    ("    3.8 Mantıksal Veritabanı Gereksinimleri", "22"),
    ("    3.9 Diğer Gereksinimler", "22"),
    ("4. Analiz Modelleri", "23"),
    ("5. Değişiklik Yönetimi Süreci", "24"),
    ("A. Ekler", "25"),
]
for label, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.font.size = Pt(11)
    run = p.add_run("\t" + page)
    run.font.size = Pt(11)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15))
page_break(doc)

# ============================================================
# 1. GİRİŞ
# ============================================================
add_heading(doc, "1. Giriş", level=1)
add_para(
    doc,
    "Bu doküman, Restoran POS v1.1.0 yazılımı için Yazılım Gereksinimleri Belirtimi (SRS) "
    "raporudur. IEEE 830-1998 standardına dayalı olarak hazırlanmış, sistemin işlevsel ve "
    "işlevsel olmayan gereksinimlerini, tasarım kısıtlamalarını ve analiz modellerini içerir.",
)

# 1.1 Amaç
add_heading(doc, "1.1 Amaç", level=2)
add_para(
    doc,
    "Bu Yazılım Gereksinimleri Belirtimi dokümanının amacı, geliştirilen Restoran POS "
    "(Point-of-Sale) sisteminin tüm fonksiyonel ve fonksiyonel olmayan gereksinimlerini "
    "açık, ölçülebilir ve doğrulanabilir biçimde belirlemektir. Doküman; yazılım geliştirici, "
    "test mühendisi, proje danışmanı ve akademik değerlendirme komisyonu tarafından "
    "okunmak üzere hazırlanmıştır.",
)
add_para(
    doc,
    "Doküman özellikle aşağıdaki paydaşlara hitap eder:",
)
add_bullet(doc, "Akademik değerlendirme jürisi — sistemin kapsamını, mimarisini ve karşılanan gereksinimleri inceler.")
add_bullet(doc, "Proje danışmanı — gereksinim-tasarım uyumunu doğrular.")
add_bullet(doc, "Geliştirici (yazar) — referans olarak gelecek sürümlerde kullanır.")
add_bullet(doc, "Pilot işletme yöneticileri — sistemin yapabileceklerini ve sınırlarını anlamak için.")

# 1.2 Kapsam
add_heading(doc, "1.2 Kapsam", level=2)
add_para(doc, "Üretilen yazılım ürünleri:", bold=True)
add_bullet(doc, "Restoran POS Desktop — Windows için Electron tabanlı masaüstü uygulaması (Restoran POS.exe).")
add_bullet(doc, "POS API Server — Yerel olarak çalışan Express tabanlı REST + Socket.io sunucusu.")
add_bullet(doc, "Store Bridge — Yerel donanım entegrasyon katmanı (ESC/POS yazıcılar, CallerID).")
add_bullet(doc, "CallerID SDK Helper — C# .NET 8 self-contained binary (HID gelen arama dinleyici).")
add_bullet(doc, "POS SQLite Veritabanı — better-sqlite3 ile WAL kipinde tek dosya saklama.")

add_para(doc, "Sistemin yapacakları:", bold=True)
add_bullet(doc, "Restoran masalarının yönetimi, sipariş alma, mutfak yönlendirme, ödeme alma.")
add_bullet(doc, "Paket sipariş ve müşteri yönetimi, çoklu telefon ve adres kayıtları.")
add_bullet(doc, "ESC/POS protokolüne uyumlu termal yazıcılarda Türkçe karakter destekli fiş basımı (PC857).")
add_bullet(doc, "Telefonla gelen aramaların CallerID donanımı üzerinden otomatik müşteri tanıması.")
add_bullet(doc, "Gerçek zamanlı (Socket.io) mutfak, masa ve paket ekranı senkronizasyonu.")
add_bullet(doc, "Rol bazlı yetkilendirme (Yönetici, Kasiyer, Garson, Mutfak).")
add_bullet(doc, "Günlük satış, ödeme dökümü, ürün ve kategori bazlı rapor ile 4 etkileşimli grafik.")
add_bullet(doc, "Otomatik DB yedekleme (gece 02:00) ve manuel geri yükleme (SHA-256 doğrulamalı).")

add_para(doc, "Sistemin yapmayacakları (kapsam dışı):", bold=True)
add_bullet(doc, "Ödeme terminal (POS cihazı) ile doğrudan entegrasyon — sağlayıcı kararı bekleniyor.")
add_bullet(doc, "Resmi e-belge / e-arşiv fatura entegrasyonu — vergi mevzuatı uyumlanmadı.")
add_bullet(doc, "Mobil garson uygulaması — ayrı bir proje kapsamında planlanmıştır.")
add_bullet(doc, "Online sipariş platformları (Yemeksepeti, Getir) entegrasyonu.")
add_bullet(doc, "Çoklu şube/tenant tabanlı SaaS özellikleri.")

add_para(doc, "Hedefler:", bold=True)
add_bullet(doc, "Tek bilgisayar üzerinde tam fonksiyonel restoran yönetimi sağlamak.")
add_bullet(doc, "Çevrimdışı (internetsiz) çalışabilen, hızlı ve güvenilir kullanıcı deneyimi sunmak.")
add_bullet(doc, "Ticari POS sistemlerine (SambaPOS, Adisyo) açık kaynak temelli alternatif geliştirmek.")
add_bullet(doc, "Türkçe restoran sektörüne özgü ihtiyaçları (PC857 yazıcı, CallerID) karşılamak.")
add_bullet(doc, "Akademik bir bitirme projesi olarak yazılım mühendisliği prensiplerini uygulamak.")

# 1.3 Tanımlar ve Kısaltmalar
add_heading(doc, "1.3 Tanımlar ve Kısaltmalar", level=2)
add_table(
    doc,
    ["Terim / Kısaltma", "Açıklama"],
    [
        ["POS", "Point-of-Sale — Satış Noktası Sistemi."],
        ["SRS", "Software Requirements Specification — Yazılım Gereksinimleri Belirtimi."],
        ["IEEE 830", "IEEE Recommended Practice for Software Requirements Specifications (1998)."],
        ["ESC/POS", "Epson Standard Code for POS — termal yazıcılar için komut seti standardı."],
        ["PC857", "DOS Türkçe karakter kodlaması; ESC/POS yazıcılarda Türkçe gösterim için kullanılır."],
        ["JWT", "JSON Web Token — kimlik doğrulama için imzalı taşıyıcı belirteç."],
        ["bcrypt", "Şifre özetleme (hashing) algoritması; cost factor ile maliyet ayarlanır."],
        ["CORS", "Cross-Origin Resource Sharing — kaynak paylaşım politikası."],
        ["CSRF", "Cross-Site Request Forgery — siteler arası istek sahteciliği."],
        ["SQLite WAL", "Write-Ahead Logging — eşzamanlı okuma ve yazmayı destekleyen SQLite kipi."],
        ["Migration", "Veritabanı şema sürüm değişikliği (numaralandırılmış, ileri-yönlü)."],
        ["CallerID", "Gelen arama numarası gösterimi; HID üzerinden POS'a aktarılır."],
        ["Electron", "Chromium + Node.js ile masaüstü uygulama çatısı."],
        ["Vite", "Modern frontend build aracı; HMR ve hızlı geliştirme döngüsü sağlar."],
        ["Sentry", "Hata izleme ve performans gözlemlenebilirliği platformu."],
        ["Pino", "Yüksek performanslı JSON loglama kütüphanesi (NDJSON çıktı)."],
        ["Socket.io", "WebSocket tabanlı çift yönlü gerçek zamanlı iletişim kütüphanesi."],
        ["NDJSON", "Newline-Delimited JSON — her satırı tek bir JSON nesnesi olan log formatı."],
        ["RBAC", "Role-Based Access Control — Rol tabanlı erişim kontrolü."],
        ["FK", "Foreign Key — Yabancı anahtar (veritabanı ilişki kısıtlaması)."],
        ["MTBF", "Mean Time Between Failures — Arızalar arası ortalama süre."],
        ["RTO / RPO", "Recovery Time / Point Objective — geri kazanım süresi / noktası hedefleri."],
        ["SmartScreen", "Windows'un imzalanmamış uygulamalar için gösterdiği güvenlik uyarısı."],
        ["NSIS", "Nullsoft Scriptable Install System — Windows kurulum dosyası üreticisi."],
    ],
    col_widths=[Cm(4), Cm(12.5)],
)

# 1.4 Referanslar
add_heading(doc, "1.4 Referanslar", level=2)
references = [
    "IEEE Std 830-1998 — IEEE Recommended Practice for Software Requirements Specifications. IEEE Computer Society, 1998.",
    "Electron Documentation — https://www.electronjs.org/docs/latest/ (Erişim: 2026-04-15).",
    "Express.js Documentation — https://expressjs.com/ (Erişim: 2026-04-15).",
    "better-sqlite3 — https://github.com/WiseLibs/better-sqlite3 (Erişim: 2026-04-15).",
    "Socket.IO Documentation — https://socket.io/docs/v4/ (Erişim: 2026-04-15).",
    "Sentry Node.js SDK — https://docs.sentry.io/platforms/javascript/guides/node/ (Erişim: 2026-04-20).",
    "Pino Logger — https://getpino.io/ (Erişim: 2026-04-20).",
    "Vitest Testing Framework — https://vitest.dev/ (Erişim: 2026-04-25).",
    "Playwright End-to-End Testing — https://playwright.dev/ (Erişim: 2026-04-25).",
    "ESC/POS Command Reference — Epson Technical Reference Guide (TM-T88V serisi).",
    "T.C. Hazine ve Maliye Bakanlığı — Yazar Kasa POS Mevzuatı (referans, kapsam dışı).",
    "OWASP Top 10 — https://owasp.org/Top10/ (Erişim: 2026-04-30).",
    "Proje Deposu — https://github.com/ilhanavc/restoran-pos",
]
for i, r in enumerate(references, start=1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] ")
    run.bold = True
    p.add_run(r)

# 1.5 Genel Bakış
add_heading(doc, "1.5 Genel Bakış", level=2)
add_para(
    doc,
    "Bu dokümanın geri kalanı IEEE 830-1998 standardının önerdiği yapıda düzenlenmiştir. "
    "İkinci bölüm sistemin genel açıklamasını, üçüncü bölüm fonksiyonel ve fonksiyonel "
    "olmayan tüm özel gereksinimleri ayrıntılı olarak verir. Dördüncü bölümde analiz "
    "modelleri (aktivite, sequence, veri akış, durum geçişi diyagramları), beşinci bölümde "
    "değişiklik yönetimi süreci ele alınır. Ek olarak son bölümde teknoloji yığını ve "
    "kurulum talimatları verilir.",
)
page_break(doc)

# ============================================================
# 2. GENEL AÇIKLAMA
# ============================================================
add_heading(doc, "2. Genel Açıklama", level=1)
add_para(
    doc,
    "Bu bölüm, Restoran POS v1.1.0 yazılımının çevresel bağlamını, ana işlevlerini, "
    "kullanıcı profilini, sınırlamalarını ve varsayımlarını yüksek seviyede açıklar. "
    "Detaylı gereksinim listesi 3. bölümde verilmiştir.",
)

# 2.1 Ürün Perspektifi
add_heading(doc, "2.1 Ürün Perspektifi", level=2)
add_para(
    doc,
    "Restoran POS, masaüstü ortamda çalışan bağımsız (self-hosted) bir POS yazılımıdır. "
    "Ticari muadilleri (SambaPOS, Adisyo, Mikrocell) genellikle aylık abonelikli ve "
    "kapalı kaynak iken bu yazılım, açık mimari, modüler ve uyarlanabilir bir alternatif "
    "olarak konumlanır. Sistem üç ana bileşenden oluşur:",
)
add_bullet(doc, "Electron tabanlı masaüstü kapsayıcı (kullanıcı arayüzü için Chromium görüntüleyici).")
add_bullet(doc, "Yerel Express API sunucusu (iş kuralları + SQLite veri katmanı + Socket.io).")
add_bullet(doc, "Store Bridge yerel donanım katmanı (yazıcı kuyruğu + CallerID).")
add_para(
    doc,
    "Tüm bileşenler aynı Windows makinesinde çalışır, internet bağlantısı gerekmez. "
    "Veriler kullanıcı profili altındaki %APPDATA%\\restoran-pos klasöründe SQLite "
    "(better-sqlite3, WAL kipi) ile saklanır. Sistem, ek olarak Sentry ile uzaktan "
    "hata gözetimi yapabilir (opsiyonel).",
)

# Mimari diyagramı tablo halinde
add_para(doc, "Sistemin yüksek seviye mimari görünümü aşağıdaki gibidir:", italic=True)
add_table(
    doc,
    ["Katman", "Bileşen", "Sorumluluk"],
    [
        ["Sunum", "React 18 + Vite (HashRouter)", "Kullanıcı arayüzü, durum yönetimi, form doğrulama"],
        ["Uygulama", "Express 4 REST API", "Auth, sipariş, ödeme, rapor; Zod doğrulama; rate limit"],
        ["Gerçek-zamanlı", "Socket.io 4", "Mutfak / masa / paket ekranlarına anlık olay yayını"],
        ["Veri", "better-sqlite3 (WAL)", "Senkron ACID; 13 numaralı migration disiplini"],
        ["Donanım", "Store Bridge", "ESC/POS yazıcı kuyruğu, PC857 Türkçe encoding"],
        ["Donanım", "CallerID Helper (.NET 8)", "HID telefon olaylarını POS'a iletme"],
        ["Paketleme", "Electron 34 + electron-builder", "Windows tek-tıkla çalışan .exe"],
    ],
    col_widths=[Cm(3), Cm(5), Cm(8.5)],
)

# 2.2 Ürün İşlevleri
add_heading(doc, "2.2 Ürün İşlevleri", level=2)
add_para(doc, "Sistemin sağladığı temel işlev grupları aşağıda özetlenmiştir:")
add_table(
    doc,
    ["Modül", "İşlev Özeti"],
    [
        ["Kimlik Doğrulama", "Rol bazlı giriş (Yönetici/Kasiyer/Garson/Mutfak); JWT + refresh token; bcrypt; zorunlu şifre değişikliği akışı; şifre politikası (min 8 karakter, büyük harf, rakam)."],
        ["Masa Yönetimi", "Alan tabanlı grid, masa doluluk renk skalası, masa transferi, kapasite tanımı."],
        ["Sipariş Akışı", "Kategori → ürün hiyerarşik seçimi, modifier (acılı/az pişmiş vb.), kalem notu, sepet, mutfağa gönderme."],
        ["Paket Sipariş", "Müşteri arama (telefon/ad), çoklu adres yönetimi, kart/nakit ödeme tipi seçim zorunluluğu."],
        ["Ödeme", "Nakit, kart, karışık; hızlı tutar butonları (50/100/200/500 TL); indirim; para üstü; otomatik kapatma."],
        ["Mutfak Ekranı", "Aktif sipariş listesi; kalem bazlı 'hazırlandı' işareti; yaş uyarıları (10dk sarı, 20dk kırmızı); Socket.io ile anlık güncelleme."],
        ["Yazıcı / Fiş", "ESC/POS PC857 encoding; kelime kaydırma; 4 şablon (paket-kasa/paket-mutfak/masa-mutfak/masa-kasa); özelleştirilebilir başlık-altlık."],
        ["Rezervasyon", "Takvim görünümü; tarih + kişi sayısı + not; masa oturtma bağlantısı; geldi / gelmedi durumları."],
        ["Stok Takibi", "Kalem CRUD; stok hareketleri; düşük stok uyarısı."],
        ["Müşteri Yönetimi", "Çoklu telefon (normalize edilmiş) ve adres; sipariş geçmişi; Excel/CSV import-export."],
        ["Raporlar", "Günlük satış, ödeme dökümü, kategori/ürün/kullanıcı dağılımı; 4 etkileşimli grafik (recharts); X/Z dönem kapatma."],
        ["İade / Bahşiş", "Sipariş ve ödemeye bağlı iade; kapalı dönem kontrolü; ödeme bazlı bahşiş; raporlarda ayrı toplam."],
        ["CallerID", "C812A V8 HID + .NET 8 SDK helper veya pano dinleme yedek yolu; gelen arama modali."],
        ["Yedek / Geri Yükle", "Gece 02:00 otomatik yedek; manuel yedek; SHA-256 doğrulama; uploads + pos-config.json dahil; iki adımlı geri yükleme modali."],
        ["Sistem Yönetimi", "Setup wizard (4 adım); yazıcı tanımları; printer routing; audit log görüntüleyici; destek paketi (support-bundle)."],
    ],
    col_widths=[Cm(4), Cm(12.5)],
)

# 2.3 Kullanıcı Özellikleri
add_heading(doc, "2.3 Kullanıcı Özellikleri", level=2)
add_para(
    doc,
    "Sistemin nihai kullanıcıları aşağıdaki rollerden oluşur. Her rol farklı yetki "
    "kümesine sahiptir ve farklı işlevsel ihtiyaçlara yöneliktir.",
)
add_table(
    doc,
    ["Rol", "Tipik Profil", "Yetki Kapsamı"],
    [
        [
            "Yönetici (Admin)",
            "Restoran sahibi / işletme müdürü. Orta-yüksek dijital okuryazarlık.",
            "Tüm modüllere erişim; menü tanımları; raporlar; yedekleme; yazıcı yapılandırması; kullanıcı yönetimi.",
        ],
        [
            "Kasiyer (Cashier)",
            "Genelde kasa başında sürekli çalışan personel. Hız ve doğruluk öncelikli.",
            "Sipariş alma, ödeme, paket sipariş, müşteri kaydı, basit rapor.",
        ],
        [
            "Garson (Waiter)",
            "Müşteri masalarına servis yapan personel. Sıklıkla mobil/tablete benzer hız ihtiyacı.",
            "Masa açma; ürün ekleme; mutfağa gönderme; garson çağrısı yanıtlama.",
        ],
        [
            "Mutfak (Kitchen)",
            "Mutfak personeli. Ekran bazlı sipariş takip ihtiyacı.",
            "Mutfak ekranı görüntüleme; kalem bazlı 'hazırlandı' işareti; sipariş durumu güncelleme.",
        ],
    ],
    col_widths=[Cm(3.5), Cm(6.5), Cm(6.5)],
)
add_para(doc, "Kullanıcı genel beklentileri:", italic=True)
add_bullet(doc, "İlk eğitimden sonra 30 dakika içinde temel akışları (masa açma, ödeme alma) yapabilir hale gelmek.")
add_bullet(doc, "Yoğun servis saatlerinde (örn: 19:00-21:00) yavaşlama yaşamamak.")
add_bullet(doc, "Sistem çökmelerinde veri kaybı yaşamamak (otomatik yedekleme + WAL kipi).")
add_bullet(doc, "Türkçe karakterleri yazıcı ve ekranlarda doğru görmek.")

# 2.4 Genel Sınırlamalar
add_heading(doc, "2.4 Genel Sınırlamalar", level=2)
add_bullet(doc, "Yalnızca Windows 10/11 (64-bit) işletim sistemlerinde çalışır; macOS ve Linux desteklenmez.")
add_bullet(doc, "Tek bilgisayar kurulumudur; çoklu kasa/şube ağ senkronizasyonu bu sürümde yoktur (v2 yol haritasında).")
add_bullet(doc, "Belge boyutu olarak en fazla 100 GB SQLite veritabanı pratikte desteklenir (better-sqlite3 sınırı).")
add_bullet(doc, "Yazıcılar ESC/POS uyumlu olmalıdır (yaygın termal yazıcı standartı).")
add_bullet(doc, "CallerID için C812A V8 veya HID uyumlu donanım gereklidir; alternatif olarak pano dinleme PowerShell scripti çalışır.")
add_bullet(doc, "Kod imzası sertifikası bulunmadığı için ilk kurulumda Windows SmartScreen uyarısı gösterilir; kullanıcı 'Yine de çalıştır' diyerek devam etmelidir.")
add_bullet(doc, "Resmi vergi belgesi (e-fatura, e-arşiv) entegrasyonu yoktur — fiş yazdırma bilgilendirme amaçlıdır.")

# 2.5 Varsayımlar ve Bağımlılıkları
add_heading(doc, "2.5 Varsayımlar ve Bağımlılıkları", level=2)
add_para(doc, "Aşağıdaki varsayımlar değişirse gereksinimler revize edilmelidir:")
add_bullet(doc, "Donanım: 64-bit Intel/AMD CPU, en az 4 GB RAM, 2 GB boş disk alanı.")
add_bullet(doc, "İşletim sistemi: Windows 10 sürüm 1809 (Build 17763) veya üstü, ya da Windows 11.")
add_bullet(doc, "Yazıcılar TCP/IP veya USB üzerinden bağlanır; LAN'da kullanılıyorsa sabit IP atanmıştır.")
add_bullet(doc, "Restoran personeli Windows GUI ile etkileşim deneyimine sahiptir (temel mouse/dokunmatik kullanımı).")
add_bullet(doc, "Sistem yöneticisi rolü gerektiğinde GitHub Releases üzerinden güncel sürümü indirebilir.")
add_bullet(doc, "Sentry uzak hata izleme opsiyoneldir; etkinleştirildiğinde internet bağlantısı şarttır ancak temel çalışma çevrimdışı sürebilir.")
add_bullet(doc, "Pilot kurulum öncesinde admin@demo.com hesabıyla seed verisi ile demo yapılabilir; canlı kuruluma geçilirken yeni admin hesabı oluşturulur.")
page_break(doc)

# ============================================================
# 3. ÖZEL GEREKSİNİMLER
# ============================================================
add_heading(doc, "3. Özel Gereksinimler", level=1)
add_para(
    doc,
    "Bu bölüm, sistemin geliştirme, test ve doğrulama aşamalarını yönlendiren özel "
    "gereksinimleri tanımlar. Tüm gereksinimler doğru, izlenebilir, kesin anlamlı, "
    "test edilebilir ve önceliklendirilmiş olacak şekilde verilmiştir. Her gereksinim "
    "kod (GR-x.x.x) ile tek başına belirlenebilir.",
)

# 3.1 Dış Arabirim Gereksinimleri
add_heading(doc, "3.1 Dış Arabirim Gereksinimleri", level=2)

add_heading(doc, "3.1.1 Kullanıcı Arayüzleri", level=3)
add_bullet(doc, "GR-3.1.1.1 — Tüm ekranlar 1366×768 ve daha yüksek çözünürlüklerde düzgün görüntülenmelidir.")
add_bullet(doc, "GR-3.1.1.2 — Arayüz dili Türkçe olarak sunulmalıdır.")
add_bullet(doc, "GR-3.1.1.3 — Hover, focus, disabled durumları görsel olarak ayırt edilebilir olmalıdır (WCAG 2.1 AA).")
add_bullet(doc, "GR-3.1.1.4 — Onay ve hata mesajları toast bileşeni ile en az 3 saniye gösterilmelidir.")
add_bullet(doc, "GR-3.1.1.5 — Yıkıcı işlemler (silme, geri yükleme) için iki adımlı onay modali kullanılmalıdır; window.confirm yerine ConfirmDialog bileşeni zorunludur.")
add_bullet(doc, "GR-3.1.1.6 — Beyaz ekran (renderer çökmesi) durumunda ErrorBoundary 'Yenile' butonu içeren güvenli ekran göstermelidir.")

add_heading(doc, "3.1.2 Donanım Arabirimleri", level=3)
add_bullet(doc, "GR-3.1.2.1 — ESC/POS uyumlu termal yazıcılar TCP (port 9100) veya USB üzerinden bağlanır; her yazıcı için ayrı kuyruk açılır.")
add_bullet(doc, "GR-3.1.2.2 — Türkçe karakterler PC857 kod sayfasıyla kodlanır; ESC t 12 komutu basım öncesi gönderilir.")
add_bullet(doc, "GR-3.1.2.3 — Yazıcı tipi başına 'skipInit' bayrağı, ESC @ komutunun gönderilmemesini seçeneğe bağlar.")
add_bullet(doc, "GR-3.1.2.4 — CallerID için HID uyumlu USB cihaz (C812A V8) ya da pano (clipboard) dinleme PowerShell scripti kullanılır.")
add_bullet(doc, "GR-3.1.2.5 — Gelen arama POST /api/bridge/caller-id/incoming uç noktasına X-Bridge-Token başlığıyla iletilir.")

add_heading(doc, "3.1.3 Yazılım Arabirimleri", level=3)
add_bullet(doc, "GR-3.1.3.1 — Frontend ile backend HTTP/JSON üzerinden (REST) haberleşir; tüm uç noktalar /api/ önekiyle başlar.")
add_bullet(doc, "GR-3.1.3.2 — Gerçek zamanlı bildirimler Socket.io üzerinden 'kitchen', 'tables', 'takeaway' kanallarında yayınlanır.")
add_bullet(doc, "GR-3.1.3.3 — Veritabanı better-sqlite3 sürücüsüyle WAL kipinde tutulur; ON DELETE / ON UPDATE FK kısıtlamaları aktiftir.")
add_bullet(doc, "GR-3.1.3.4 — Store Bridge yerel HTTP üzerinden Bridge Token doğrulamasıyla print_jobs kuyruğunu çeker.")
add_bullet(doc, "GR-3.1.3.5 — Otomatik güncelleme electron-updater + GitHub Releases üzerinden kontrol edilir.")

add_heading(doc, "3.1.4 İletişim Arabirimleri", level=3)
add_bullet(doc, "GR-3.1.4.1 — Tüm dış HTTP istekleri CORS whitelist kontrolünden geçer; dev ortamı için localhost / 192.168.x / 10.x bilgisi otomatik kabul edilir.")
add_bullet(doc, "GR-3.1.4.2 — Sentry uzak servisi yalnızca yapılandırıldığında çağrılır; aksi halde tüm hata izleme yereldedir.")
add_bullet(doc, "GR-3.1.4.3 — electron-updater GitHub API'sini 10 dakika aralıkla sorgular; yeni sürüm bulunursa kullanıcıya bildirim gösterir.")

# 3.2 Fonksiyonel Gereksinimler
add_heading(doc, "3.2 Fonksiyonel Gereksinimleri", level=2)
add_para(
    doc,
    "Bu bölüm sistemde uygulanan ana fonksiyonel gereksinimleri her biri Giriş, "
    "Girişleri, İşleme, Çıkışları ve Hata İşleme alt başlıklarıyla detaylandırır.",
)

# 3.2.1 — Kimlik Doğrulama
add_heading(doc, "3.2.1 Kimlik Doğrulama ve Yetkilendirme", level=3)
add_para(doc, "3.2.1.1 Giriş", bold=True)
add_para(
    doc,
    "Kullanıcı, e-posta ve şifre kombinasyonu ile sisteme giriş yapar. Sistem JWT "
    "access token (1 saat) ve mobil cihazlar için refresh token (30 gün) üretir. "
    "Erişim kontrolü rol tabanlıdır (RBAC).",
)
add_para(doc, "3.2.1.2 Girişleri", bold=True)
add_bullet(doc, "Email (string, RFC 5322 formatı doğrulanır)")
add_bullet(doc, "Şifre (string, minimum 8 karakter politikası yalnızca yeni hesap/şifre değişimi sırasında zorunludur)")
add_para(doc, "3.2.1.3 İşleme", bold=True)
add_bullet(doc, "Email + business_id kombinasyonu users tablosundan sorgulanır.")
add_bullet(doc, "bcrypt.compare ile şifre hash'i doğrulanır (cost factor 10).")
add_bullet(doc, "users.must_change_password=1 ise 403 ile yönlendirme zorunludur.")
add_bullet(doc, "Başarılı doğrulamada JWT imzalanır, refresh token üretilir (mobil için).")
add_bullet(doc, "Başarısız login audit_logs tablosuna 'login_failed' kaydı atar.")
add_para(doc, "3.2.1.4 Çıkışları", bold=True)
add_bullet(doc, "HTTP 200: { token, user, businessId, role }")
add_bullet(doc, "HTTP 401: { error: 'Geçersiz e-posta veya şifre' } (genel mesaj — kullanıcı keşfini önlemek için)")
add_bullet(doc, "HTTP 403: { must_change_password: true, email, businessId } (zorunlu şifre değişim akışı)")
add_para(doc, "3.2.1.5 Hata İşleme", bold=True)
add_bullet(doc, "5 başarısız denemeden sonra 15 dakika boyunca aynı IP'den login'e rate-limit (5/15dk) uygulanır.")
add_bullet(doc, "Hatalar Sentry'ye gönderilir; ancak email/şifre alanları SENSITIVE_KEYS redaksiyonu ile maskelenir.")

# 3.2.2 — Sipariş Akışı
add_heading(doc, "3.2.2 Masa Sipariş Yönetimi", level=3)
add_para(doc, "3.2.2.1 Giriş", bold=True)
add_para(
    doc,
    "Yetkili kullanıcı (Yönetici/Kasiyer/Garson) bir masa seçerek sipariş açar, kategori "
    "ve ürünleri sepete ekler, kalem notu girer ve siparişi mutfağa gönderir. "
    "Mutfak ekranı Socket.io üzerinden anlık güncellenir.",
)
add_para(doc, "3.2.2.2 Girişleri", bold=True)
add_bullet(doc, "table_id (UUID)")
add_bullet(doc, "items: [{ product_id, quantity, note?, modifiers[] }]")
add_bullet(doc, "service_charge_rate (opsiyonel, snapshot'lanır)")
add_para(doc, "3.2.2.3 İşleme", bold=True)
add_bullet(doc, "POST /api/orders → orderService.createOrder")
add_bullet(doc, "Veritabanı transaction içinde: orders, order_items, entity_mutations kayıtları yazılır.")
add_bullet(doc, "Her order_item için vat_rate_snapshot ve ürün adı snapshot kayıtları yazılır (gelecek menü değişikliklerine karşı).")
add_bullet(doc, "Tutarlar hem REAL (TL) hem _cents (integer minor unit) olarak yazılır.")
add_bullet(doc, "Socket.io 'order:created' eventi mutfak kanalına yayınlanır.")
add_bullet(doc, "Print queue'ya mutfak fişi job'u eklenir (printer routing'e göre).")
add_para(doc, "3.2.2.4 Çıkışları", bold=True)
add_bullet(doc, "HTTP 201: yeni sipariş JSON nesnesi (id, status='active', grand_total, items[])")
add_bullet(doc, "Socket: kitchen kanalına gerçek zamanlı bildirim")
add_bullet(doc, "Yazıcıda mutfak fişi (kategori bazlı routing ile)")
add_para(doc, "3.2.2.5 Hata İşleme", bold=True)
add_bullet(doc, "Geçersiz product_id → Zod validation 400.")
add_bullet(doc, "Masa zaten 'paid' → 409 Conflict.")
add_bullet(doc, "Yazıcı erişilemezse sipariş başarılı sayılır ancak print_jobs 'failed' işaretlenir; admin manuel retry edebilir.")

# 3.2.3 — Ödeme
add_heading(doc, "3.2.3 Ödeme Alma", level=3)
add_para(doc, "3.2.3.1 Giriş", bold=True)
add_para(
    doc,
    "Açık bir sipariş üzerinden nakit, kart veya karışık tutarlarla ödeme alınır. "
    "İndirim uygulanabilir; para üstü otomatik hesaplanır. Tam ödeme tamamlandığında "
    "sipariş 'paid' durumuna geçer ve masa boşa düşer.",
)
add_para(doc, "3.2.3.2 Girişleri", bold=True)
add_bullet(doc, "order_id (UUID)")
add_bullet(doc, "payments: [{ method: 'cash'|'card', amount, tip? }]")
add_bullet(doc, "discount_amount (opsiyonel)")
add_para(doc, "3.2.3.3 İşleme", bold=True)
add_bullet(doc, "POST /api/payments → paymentService.recordPayment")
add_bullet(doc, "Transaction içinde payments, refunds, entity_mutations yazılır.")
add_bullet(doc, "Toplam ödenen >= grand_total ise order.status = 'paid' yapılır.")
add_bullet(doc, "Kapalı dönem (period close) sonrası iade yapılamaz; guard kontrolü vardır.")
add_bullet(doc, "Kasa fişi print queue'ya eklenir.")
add_para(doc, "3.2.3.4 Çıkışları", bold=True)
add_bullet(doc, "HTTP 201: { payment, order: { status, balance, change } }")
add_bullet(doc, "Kasa fişi yazıcıda")
add_bullet(doc, "Masa durum güncellemesi Socket.io ile yayınlanır")
add_para(doc, "3.2.3.5 Hata İşleme", bold=True)
add_bullet(doc, "Negatif amount → 400 Bad Request.")
add_bullet(doc, "Ödeme tutarı sipariş bakiyesinden büyükse 'overpayment' durumunda izin verilir (para üstü hesaplanır).")
add_bullet(doc, "Sipariş zaten paid ise 409 Conflict.")

# 3.2.4 — Yazıcı Kuyruğu
add_heading(doc, "3.2.4 Yazıcı Kuyruğu ve Fiş Basımı", level=3)
add_para(doc, "3.2.4.1 Giriş", bold=True)
add_para(
    doc,
    "Sipariş ve ödeme olaylarında otomatik üretilen yazıcı işleri, lease-tabanlı bir "
    "kuyruk üzerinden Store Bridge tarafından alınır. ESC/POS PC857 kodlamasıyla "
    "Türkçe karakter destekli fiş üretilir.",
)
add_para(doc, "3.2.4.2 Girişleri", bold=True)
add_bullet(doc, "print_jobs.status='pending' satırlar (poller her 2 saniyede sorgular)")
add_bullet(doc, "Lease süresi: 60 saniye")
add_para(doc, "3.2.4.3 İşleme", bold=True)
add_bullet(doc, "Bridge claim: claimed_until güncellenir, bridge_id atanır.")
add_bullet(doc, "Renderer (4 şablon): paket-kasa / paket-mutfak / masa-mutfak / masa-kasa.")
add_bullet(doc, "32 karakter genişliği için kelime kaydırma (word-wrap) uygulanır.")
add_bullet(doc, "TCP veya USB üzerinden ESC/POS verisi gönderilir.")
add_bullet(doc, "Başarılı basımda status='completed'; başarısızda 'failed' + last_error_code.")
add_para(doc, "3.2.4.4 Çıkışları", bold=True)
add_bullet(doc, "Termal yazıcıda Türkçe karakterli fiş")
add_bullet(doc, "print_jobs.completed_at timestamp güncellenir")
add_para(doc, "3.2.4.5 Hata İşleme", bold=True)
add_bullet(doc, "16 ESC/POS hata kodu dictionary'si (Türkçe label + öneri): tcp_timeout, encoding_error, paper_out, vb.")
add_bullet(doc, "Failed iş için admin manuel retry edebilir; otomatik retry yoktur (mükerrer mutfak fişi engeli).")

# 3.2.5 — Yedek/Geri Yükleme
add_heading(doc, "3.2.5 Yedek Alma ve Geri Yükleme", level=3)
add_para(doc, "3.2.5.1 Giriş", bold=True)
add_para(
    doc,
    "Sistem günlük olarak gece 02:00'de otomatik yedek alır; admin manuel olarak da "
    "tetikleyebilir. Yedekler SHA-256 hash ile imzalanır ve meta.json sidecar dosyası "
    "ile birlikte saklanır.",
)
add_para(doc, "3.2.5.2 Girişleri", bold=True)
add_bullet(doc, "Tetikleyici: cron (02:00), manuel butona tıklama, harici robocopy task")
add_bullet(doc, "Veri kaynağı: %APPDATA%\\restoran-pos\\pos.db + uploads + pos-config.json")
add_para(doc, "3.2.5.3 İşleme", bold=True)
add_bullet(doc, "WAL checkpoint TRUNCATE; tek dosya halinde DB snapshot.")
add_bullet(doc, "uploads/products/ klasörü dahil edilir.")
add_bullet(doc, "meta.json: appVersion, schemaVersion, rowCounts, integrityCheck, sha256.")
add_bullet(doc, "30 gün tutma (rotation); eskiyen yedekler silinir.")
add_para(doc, "3.2.5.4 Çıkışları", bold=True)
add_bullet(doc, "backups/backup-YYYYMMDD-HHMMSS.zip + meta.json")
add_bullet(doc, "Toast bildirimi (manuel yedek için)")
add_para(doc, "3.2.5.5 Hata İşleme", bold=True)
add_bullet(doc, "Disk alanı yetersizse ön kontrol uyarısı verilir.")
add_bullet(doc, "Geri yükleme sonrası bütünlük kontrolü başarısız olursa otomatik safety revert.")
add_bullet(doc, "backup-failed IPC kanalı App.jsx → toast.warning ile kullanıcıya bildirilir.")

# 3.3 Kullanım Durumları
add_heading(doc, "3.3 Kullanım Durumları (Use Cases)", level=2)

add_heading(doc, "3.3.1 Kullanım Durumu #1 — Masa Sipariş ve Ödeme Tam Akış", level=3)
add_table(
    doc,
    ["Alan", "İçerik"],
    [
        ["Kimlik", "UC-01"],
        ["Aktörler", "Kasiyer (birincil), Mutfak (ikincil), Yazıcı (sistem)"],
        ["Ön koşul", "Kasiyer giriş yapmış, en az 1 boş masa, en az 1 aktif ürün mevcut"],
        ["Tetikleyici", "Kasiyer Masalar ekranında bir masaya tıklar"],
        ["Ana Akış",
            "1. Kasiyer M1 masasına tıklar; sipariş ekranı açılır\n"
            "2. Kategori seçer, ürünleri sepete ekler\n"
            "3. 'Mutfağa Gönder' butonuna tıklar; sistem siparişi kaydeder, mutfak ekranı anında güncellenir\n"
            "4. Sipariş hazır olduğunda mutfak personeli 'hazırlandı' işaretler\n"
            "5. Kasiyer 'Ödeme' butonuna tıklar\n"
            "6. Nakit girilir, 'Ödemeyi Tamamla' tıklanır\n"
            "7. Para üstü görüntülenir, masa boşa düşer, kasa fişi yazıcıdan çıkar"],
        ["Alternatif Akış",
            "3a. Yazıcı erişilemiyorsa sipariş kaydedilir, kullanıcıya uyarı gösterilir, admin manuel yeniden basım yapar"],
        ["Son Durum", "Sipariş 'paid' durumunda, masa 'boş' durumunda, kasa fişi basılmıştır"],
    ],
    col_widths=[Cm(3.5), Cm(13)],
)

add_heading(doc, "3.3.2 Kullanım Durumu #2 — Paket Sipariş ve CallerID Entegrasyonu", level=3)
add_table(
    doc,
    ["Alan", "İçerik"],
    [
        ["Kimlik", "UC-02"],
        ["Aktörler", "Müşteri (telefon), Kasiyer, CallerID Helper (sistem)"],
        ["Ön koşul", "CallerID HID cihazı bağlı veya pano dinleme aktif; X-Bridge-Token tanımlı"],
        ["Tetikleyici", "Müşteri restoranın sabit hattını arar"],
        ["Ana Akış",
            "1. CallerID cihazı gelen aramayı algılar\n"
            "2. .NET 8 SDK helper veya PowerShell pano dinleyicisi numarayı POS API'ye POST eder\n"
            "3. Sistem numarayı normalize eder ve customers tablosunda arar\n"
            "4. POS ön ekranında 'Gelen Arama' modali açılır (kayıtlı müşteri varsa adı görünür)\n"
            "5. Kasiyer modal üzerinden 'Paket Sipariş Oluştur' butonuna tıklar\n"
            "6. Müşteri bilgileri ve adresi otomatik dolar\n"
            "7. Ürünler eklenir, ödeme tipi seçilir, sipariş oluşturulur"],
        ["Alternatif Akış",
            "3a. Müşteri kayıtlı değilse modal 'Yeni Müşteri' formuyla açılır\n"
            "4a. Modal hızla kapanırsa /takeaway sayfasından yeniden açılabilir"],
        ["Son Durum", "Paket sipariş oluşturulmuş, mutfak fişi basılmış, müşteri kayıt veya güncellenmiştir"],
    ],
    col_widths=[Cm(3.5), Cm(13)],
)

add_heading(doc, "3.3.3 Kullanım Durumu #3 — Yedek Alma ve Geri Yükleme", level=3)
add_table(
    doc,
    ["Alan", "İçerik"],
    [
        ["Kimlik", "UC-03"],
        ["Aktörler", "Yönetici, Zamanlanmış görev (cron 02:00), Sistem"],
        ["Ön koşul", "Disk alanı en az 200 MB serbest, restore için Setup Wizard tamamlanmış"],
        ["Tetikleyici", "Otomatik: gece 02:00; veya manuel: Bakım ve Yedekleme ekranı"],
        ["Ana Akış",
            "1. WAL checkpoint TRUNCATE yapılır\n"
            "2. uploads, pos-config.json eklenir\n"
            "3. meta.json (rowCounts + sha256 + integrityCheck) yazılır\n"
            "4. ZIP arşivi backups/ klasörüne yazılır\n"
            "5. 30 günden eski yedekler silinir (rotation)\n"
            "6. Restore için: yedek seç → iki adımlı modal → bütünlük kontrolü → safety revert pozisyonu → DB değiştir"],
        ["Alternatif Akış",
            "6a. Açık sipariş varsa restore başlatılmadan uyarı gösterilir (GET /maintenance/open-orders)\n"
            "6b. Bütünlük kontrolü fail olursa otomatik safety revert tetiklenir"],
        ["Son Durum", "Yedek başarıyla alınmış veya geri yüklenmiş, sistem tutarlıdır"],
    ],
    col_widths=[Cm(3.5), Cm(13)],
)

# 3.4 Sınıflar / Nesneler
add_heading(doc, "3.4 Sınıflar / Nesneler", level=2)

add_heading(doc, "3.4.1 Order (Sipariş)", level=3)
add_para(doc, "Öznitelikler", bold=True)
add_bullet(doc, "id: UUID — birincil anahtar")
add_bullet(doc, "business_id, branch_id, table_id, user_id: ilgili tablolara FK")
add_bullet(doc, "type: 'table' | 'takeaway'")
add_bullet(doc, "status: 'active' | 'paid' | 'cancelled' | 'closed'")
add_bullet(doc, "grand_total (REAL) + grand_total_cents (INTEGER) dual-write")
add_bullet(doc, "service_charge_rate / service_charge_amount snapshot (DB-4)")
add_bullet(doc, "pricing_policy_version: oluşturma anındaki ürün MAX updated_at (DB-4)")
add_bullet(doc, "items: OrderItem[] (1-N ilişki)")
add_bullet(doc, "payments: Payment[] (1-N ilişki)")
add_bullet(doc, "created_at, updated_at: timestamp")
add_para(doc, "İşlevleri (referans: GR-3.2.2, GR-3.2.3)", bold=True)
add_bullet(doc, "createOrder(payload) — yeni sipariş oluşturur, mutfak fişi basar")
add_bullet(doc, "addItem(orderId, item) — sepete kalem ekler")
add_bullet(doc, "voidItem(orderId, itemId) — kalem siler (orderActionPolicy kontrolü)")
add_bullet(doc, "applyDiscount(orderId, amount) — indirim uygular")
add_bullet(doc, "close(orderId) — paid durumuna geçirir")

add_heading(doc, "3.4.2 Customer (Müşteri)", level=3)
add_para(doc, "Öznitelikler", bold=True)
add_bullet(doc, "id: UUID")
add_bullet(doc, "business_id: FK")
add_bullet(doc, "full_name, total_orders")
add_bullet(doc, "phones: CustomerPhone[] (normalized_phone alanı ile tek anahtar)")
add_bullet(doc, "addresses: CustomerAddress[]")
add_bullet(doc, "first_visit, last_visit timestamp")
add_para(doc, "İşlevleri", bold=True)
add_bullet(doc, "findByPhone(normalized) — telefon ile arama")
add_bullet(doc, "addPhone(customerId, phone) — yeni telefon ekleme")
add_bullet(doc, "addAddress(customerId, address) — yeni adres ekleme")
add_bullet(doc, "importFromExcel(buffer) — toplu import")

add_heading(doc, "3.4.3 PrintJob (Yazıcı İşi)", level=3)
add_para(doc, "Öznitelikler", bold=True)
add_bullet(doc, "id: UUID")
add_bullet(doc, "printer_id: FK — printers tablosu")
add_bullet(doc, "payload_json: TEXT — render edilmiş ESC/POS komutları")
add_bullet(doc, "template_name: 'kitchen' | 'receipt' | 'takeaway-kitchen' | 'takeaway-receipt'")
add_bullet(doc, "status: 'pending' | 'claimed' | 'completed' | 'failed'")
add_bullet(doc, "claimed_until: TIMESTAMP (lease süresi)")
add_bullet(doc, "bridge_id: claim sahibi bridge UUID")
add_bullet(doc, "last_error_code: 16 koddan biri (failed durumda)")
add_bullet(doc, "attempt_count, completed_at")
add_para(doc, "İşlevleri", bold=True)
add_bullet(doc, "enqueue(payload) — yeni iş oluşturur")
add_bullet(doc, "claim(bridgeId) — lease ile sahiplenir")
add_bullet(doc, "complete(jobId) — başarılı tamamlama")
add_bullet(doc, "fail(jobId, errorCode) — hata ile sonlanma")
add_bullet(doc, "retry(jobId) — admin manuel yeniden deneme")

# 3.5 İşlevsel Olmayan Gereksinimler
add_heading(doc, "3.5 İşlevsel Olmayan Gereksinimler", level=2)

add_heading(doc, "3.5.1 Performans", level=3)
add_bullet(doc, "GR-3.5.1.1 — Login isteği p95 < 500 ms olacaktır (yerel SQLite okuma + bcrypt karşılaştırması).")
add_bullet(doc, "GR-3.5.1.2 — Sipariş oluşturma p95 < 200 ms olacaktır (transaction + Socket.io broadcast dahil).")
add_bullet(doc, "GR-3.5.1.3 — Mutfak ekranı bir sipariş yayını için p95 < 100 ms gecikmeyle güncellenecektir.")
add_bullet(doc, "GR-3.5.1.4 — Uygulama açılışı (Restoran POS.exe → login ekranı) p95 < 8 saniye olacaktır.")
add_bullet(doc, "GR-3.5.1.5 — Veritabanı 100 MB civarında dolu olduğunda raporlar p95 < 2 saniyede dönecektir.")

add_heading(doc, "3.5.2 Güvenilirlik", level=3)
add_bullet(doc, "GR-3.5.2.1 — Hedef MTBF: 30 gün (yerel deployment).")
add_bullet(doc, "GR-3.5.2.2 — Veri kaybı toleransı RPO: 24 saat (gece otomatik yedek). Manuel yedek için 0 dakika.")
add_bullet(doc, "GR-3.5.2.3 — Geri kazanım süresi RTO: 15 dakika (yedekten DB ile uploads geri yüklemesi).")
add_bullet(doc, "GR-3.5.2.4 — Store Bridge çökerse 10 saniye içinde otomatik yeniden başlatılır (max 10 deneme guard ile).")
add_bullet(doc, "GR-3.5.2.5 — Sipariş, ödeme ve baskı işlemleri atomic transaction içinde yapılır (all-or-nothing).")

add_heading(doc, "3.5.3 Kullanılabilirlik", level=3)
add_bullet(doc, "GR-3.5.3.1 — Yeni bir kullanıcı 30 dakikalık eğitimden sonra temel akışları (sipariş alma, ödeme) bağımsız yapabilir.")
add_bullet(doc, "GR-3.5.3.2 — Tüm hata mesajları Türkçe ve eylem önerisi içerir (örn: 'Yazıcıya bağlanılamadı — kabloyu kontrol edin').")
add_bullet(doc, "GR-3.5.3.3 — Tüm yıkıcı işlemler iki adımlı onay modali (ConfirmDialog) ile korunur.")
add_bullet(doc, "GR-3.5.3.4 — Setup Wizard yeni kurulumda 4 adımda tamamlanır (Hoş geldiniz, İşletme, Admin, Tamamla).")

add_heading(doc, "3.5.4 Güvenlik", level=3)
add_bullet(doc, "GR-3.5.4.1 — Şifreler bcrypt cost 10 ile hash'lenir; plain text saklama yoktur.")
add_bullet(doc, "GR-3.5.4.2 — JWT_SECRET production ortamında min 32 karakter zorunludur; fail-fast guard mevcuttur.")
add_bullet(doc, "GR-3.5.4.3 — Refresh token SHA-256 ile hash'lenip saklanır; ham token kaydedilmez.")
add_bullet(doc, "GR-3.5.4.4 — Login/admin/bridge/printer endpoint'leri için ayrı rate-limit uygulanır.")
add_bullet(doc, "GR-3.5.4.5 — Tüm endpoint'ler Zod şeması ile girdi doğrulamasından geçirilir.")
add_bullet(doc, "GR-3.5.4.6 — better-sqlite3 prepared statements kullanılır; SQL injection vektörü kapalıdır.")
add_bullet(doc, "GR-3.5.4.7 — CORS production whitelist; geliştirme localhost/LAN otomatik allow.")
add_bullet(doc, "GR-3.5.4.8 — Sentry ve Pino loglarda hassas anahtarlar (password, token, cookie, jwtSecret, bridgeToken) redact edilir.")
add_bullet(doc, "GR-3.5.4.9 — Sentry session replay maskAllText + maskAllInputs + blockAllMedia ile GDPR/KVKK uyumludur.")
add_bullet(doc, "GR-3.5.4.10 — Şifre politikası: yeni hesap/şifre değişiminde min 8 karakter + büyük harf + rakam zorunlu.")

add_heading(doc, "3.5.5 Sürdürülebilirlik", level=3)
add_bullet(doc, "GR-3.5.5.1 — ESLint warning sayısı 0 olarak korunur (lint:ci --max-warnings 0 CI kapısı).")
add_bullet(doc, "GR-3.5.5.2 — Test kapsamı: backend 430, frontend 22, e2e 2 senaryo; toplam 452.")
add_bullet(doc, "GR-3.5.5.3 — Veritabanı migration disiplini: numbered, forward-only, schema_migrations tablosu.")
add_bullet(doc, "GR-3.5.5.4 — Pino loglarda her satır geçerli JSON (NDJSON); makine okunabilir.")
add_bullet(doc, "GR-3.5.5.5 — Modüler kod ayrımı: orderService, paymentService, refundService, periodCloseService domain katmanları.")
add_bullet(doc, "GR-3.5.5.6 — Runbook dokümantasyonu: kurulum, yedekleme, kod imzası, yazıcı acceptance, Sentry setup.")

add_heading(doc, "3.5.6 Taşınabilirlik", level=3)
add_bullet(doc, "GR-3.5.6.1 — Tek dosya .exe ile taşınır; kurulum gerektirmez (portable mode).")
add_bullet(doc, "GR-3.5.6.2 — Veritabanı, yedekler ve uploads %APPDATA%\\restoran-pos altında izole edilir.")
add_bullet(doc, "GR-3.5.6.3 — Backend Express ve frontend React 18 cross-platform Node.js modülleridir; macOS/Linux paketleme ileride mümkündür.")

# 3.6 Ters Gereksinimler
add_heading(doc, "3.6 Ters Gereksinimler", level=2)
add_bullet(doc, "Sistem hiçbir koşulda kullanıcı şifresini plain text saklamayacaktır.")
add_bullet(doc, "Sistem üretim ortamında JWT_SECRET olmadan veya 32 karakterden kısa anahtarla başlamayacaktır (fail-fast).")
add_bullet(doc, "Sistem failed yazıcı işini otomatik olarak yeniden denemeyecektir — yalnızca admin manuel retry yapabilir (duplicate fiş engeli).")
add_bullet(doc, "Sistem kapalı dönemli (period close) bir tarihe iade kaydı yazamaz.")
add_bullet(doc, "Sistem geri yükleme bütünlük kontrolü başarısız olursa eski veriye otomatik dönecek; kullanıcıya bilinmeyen durum bırakılmayacaktır.")
add_bullet(doc, "Sistem audit_logs ve entity_mutations kayıtlarını silmeyecek (immutable append-only).")

# 3.7 Tasarım Kısıtlamaları
add_heading(doc, "3.7 Tasarım Kısıtlamaları", level=2)
add_bullet(doc, "Platform: Yalnızca Windows 10/11 (64-bit). Hedef Electron 34, Node.js 20+.")
add_bullet(doc, "Yazıcı: ESC/POS uyumlu termal yazıcılar (genelde 80 mm, 32 karakter).")
add_bullet(doc, "Karakter kodlaması: PC857 (DOS Türkçe) — yazıcı sürücüsünden bağımsız.")
add_bullet(doc, "Veritabanı: SQLite (tek dosya, WAL kipi). PostgreSQL veya MySQL bu sürümde kullanılmaz.")
add_bullet(doc, "Paketleme: electron-builder 24.13.3 (25.x sürümünde 7za hatası vardır, yükseltilmez).")
add_bullet(doc, "Kod imzası sertifikası bu sürümde yoktur (SmartScreen uyarısı kalır).")
add_bullet(doc, "İç piyasa odaklı: tüm UI Türkçedir; çoklu dil bu sürüm kapsamında değildir.")
add_bullet(doc, "Lisans modeli: bitirme projesi olduğu için tek bilgisayar kullanımına yöneliktir; ticari lisans dağıtımı yapılmaz.")

# 3.8 Mantıksal Veritabanı Gereksinimleri
add_heading(doc, "3.8 Mantıksal Veritabanı Gereksinimleri", level=2)
add_para(doc, "Veri saklama formatı: SQLite (tek dosya, WAL kipi). Veritabanı şeması 13 numaralı migration ile yönetilir:")
add_table(
    doc,
    ["Migration", "Açıklama"],
    [
        ["0000_baseline_legacy_schema", "Mevcut şemanın baseline noktası"],
        ["0001_create_entity_mutations", "Audit trail için entity_mutations tablosu"],
        ["0002_add_cents_columns", "Para alanları için _cents shadow kolonları"],
        ["0003_snapshot_columns", "Sipariş anındaki pricing/service/vat snapshot kolonları"],
        ["0004_refresh_tokens", "Mobil refresh token rotation desteği"],
        ["0005_devices", "Cihaz pairing ve QR token tabloları"],
        ["0006_takeaway_planned_payment_type", "Paket sipariş için planlı ödeme tipi"],
        ["0007_drop_legacy_customer_columns", "Eski müşteri kolonlarının kaldırılması"],
        ["0008_customer_name_split_and_address_admin", "Ad-soyad ayrımı + adres admin yönetimi"],
        ["0009_renormalize_customer_phones", "Telefon normalize (E.164 benzeri)"],
        ["0010_cleanup_soft_deleted_categories", "Pasif kategorilerin temizlenmesi"],
        ["0011_must_change_password", "Zorunlu şifre değişimi flag'i"],
        ["0012_password_reset_requests", "Şifre sıfırlama talepleri"],
    ],
    col_widths=[Cm(6.5), Cm(10)],
)
add_para(doc, "Ana tablolar (özet, FK ilişkileri ile):", italic=True)
add_bullet(doc, "businesses → branches → users / dining_areas / tables / categories / products / customers")
add_bullet(doc, "tables → orders → order_items → product / modifiers")
add_bullet(doc, "orders → payments / refunds")
add_bullet(doc, "customers → customer_phones / customer_addresses")
add_bullet(doc, "printers → print_jobs (lease + status); printer_routing → categories")
add_bullet(doc, "audit_logs (eylem geçmişi); entity_mutations (before/after JSON snapshot)")
add_bullet(doc, "settings (anahtar-değer JSON); schema_migrations (migration tracking)")
add_bullet(doc, "refresh_tokens (mobil); devices + device_pairing_tokens (cihaz eşleştirme)")
add_para(doc, "Veri bütünlüğü kuralları:", bold=True)
add_bullet(doc, "Tüm ID'ler UUID v4 string olarak saklanır.")
add_bullet(doc, "Para alanları hem REAL (TL) hem _cents (INTEGER) olarak yazılır (DB-3 dual-write).")
add_bullet(doc, "FK kısıtlamaları PRAGMA foreign_keys=ON ile aktiftir.")
add_bullet(doc, "Soft-delete: order_items.product_id NOT NULL olduğu için sipariş geçmişi olan ürünler is_deleted=1 ile pasifleştirilir; geçmişi olmayanlar hard-delete edilir.")

# 3.9 Diğer Gereksinimler
add_heading(doc, "3.9 Diğer Gereksinimler", level=2)
add_bullet(doc, "Loglama: Pino structured (NDJSON), Sentry uzak izleme (opsiyonel).")
add_bullet(doc, "Telemetri: Sentry'de release tag 'restoran-pos@1.1.0' ile kaynak haritası eşleştirmesi.")
add_bullet(doc, "Otomatik güncelleme: electron-updater + latest.yml manifest dosyası.")
add_bullet(doc, "Destek paketi: GET /admin/support-bundle ile sistem, DB, bridge, kuyruk, log özetleri tek tıkla indirilebilir.")
page_break(doc)

# ============================================================
# 4. ANALİZ MODELLERİ
# ============================================================
add_heading(doc, "4. Analiz Modelleri", level=1)
add_para(
    doc,
    "Bu bölüm, gereksinimlerin geliştirilmesinde kullanılan analiz modellerini özetler. "
    "Detaylı çizimler proje deposu dokümantasyonunda (docs/) bulunabilir.",
)

add_heading(doc, "4.1 Aktivite Diyagramları", level=2)
add_para(doc, "Sistemde aşağıdaki ana iş akışları aktivite diyagramı olarak modellenmiştir:")
add_bullet(doc, "Sipariş Alma → Mutfak Yönlendirme → Ödeme → Masa Kapama (UC-01)")
add_bullet(doc, "Gelen Çağrı → Müşteri Tanıma → Paket Sipariş Açma (UC-02)")
add_bullet(doc, "Otomatik Yedek (cron 02:00) → SHA-256 → 30 gün rotation (UC-03)")
add_bullet(doc, "Setup Wizard (yeni kurulum) → İşletme Tanımı → Admin Şifre → Tamamla")
add_para(doc, "İllüstrasyon: ekli görseller (Ek A.1).", italic=True)

add_heading(doc, "4.2 Sequence Diyagramları", level=2)
add_bullet(doc, "Login → JWT üretimi → Frontend store → Yönlendirme")
add_bullet(doc, "Sipariş POST → orderService.transaction → Socket.io broadcast → Mutfak güncelleme")
add_bullet(doc, "Print queue: bridge poll → claim → ESC/POS render → TCP gönder → complete/fail")
add_bullet(doc, "CallerID HID → SDK helper → POST /api/bridge/caller-id/incoming → Socket → Frontend modal")
add_bullet(doc, "Refresh token rotate → eski hash invalidate → yeni hash kaydet → response")

add_heading(doc, "4.3 Veri Akış Diyagramları", level=2)
add_bullet(doc, "Seviye 0: Restoran POS Sistemi — kullanıcı, yazıcı, telefon ağı, Sentry harici varlıkları")
add_bullet(doc, "Seviye 1: Frontend → API → SQLite + Print Queue → Store Bridge → Yazıcı")
add_bullet(doc, "Detaylı veri akışı: müşteri arama, sipariş hazırlama, ödeme alma, rapor üretimi")

add_heading(doc, "4.4 Durum Geçişi Diyagramları", level=2)
add_bullet(doc, "Sipariş durumu: active → paid / cancelled (terminal)")
add_bullet(doc, "Masa durumu: empty → occupied → empty (sipariş açılır/kapanır)")
add_bullet(doc, "Print job: pending → claimed → completed | failed → (manuel retry) → pending")
add_bullet(doc, "Backup: scheduled → running → success / failed → archived")
page_break(doc)

# ============================================================
# 5. DEĞİŞİKLİK YÖNETİMİ SÜRECİ
# ============================================================
add_heading(doc, "5. Değişiklik Yönetimi Süreci", level=1)
add_para(
    doc,
    "Proje kapsamında veya gereksinimlerde değişiklik söz konusu olduğunda aşağıdaki "
    "süreç izlenir:",
)
add_bullet(doc, "1. Değişiklik talebi GitHub Issues'a 'enhancement' veya 'bug' etiketiyle kayıt edilir.")
add_bullet(doc, "2. Etki analizi yapılır: ilgili modüller, testler, dokümantasyon tespit edilir.")
add_bullet(doc, "3. Onay verilirse 'codex/<kısa-açıklama>' adlı bir feature branch açılır.")
add_bullet(doc, "4. Geliştirme tamamlandığında ESLint warning 0 ve tüm testler yeşil olmalıdır (CI kapısı).")
add_bullet(doc, "5. Pull Request açılır, kod incelemesi yapılır (en az 1 onay).")
add_bullet(doc, "6. main branch'e merge edilir, yeni semver tag çekilir (v1.x.y).")
add_bullet(doc, "7. SRS belgesinde ilgili gereksinim güncellenir veya yeni gereksinim eklenir.")
add_bullet(doc, "8. Üretim için electron-builder ile yeni release oluşturulur, GitHub Releases'a yüklenir.")
add_bullet(doc, "9. Sahada (pilot işletme) güncelleme electron-updater ile otomatik dağıtılır.")
add_para(
    doc,
    "Acil hata düzeltmeleri için hızlı yol: hotfix/<konu> branch, hızlı test, patch sürüm "
    "(v1.x.y+1) ile yayın. Tüm değişiklikler git tarihinde izlenebilir; Sentry üzerinde "
    "release tag karşılığı kontrol edilir.",
    italic=True,
)
page_break(doc)

# ============================================================
# A. EKLER
# ============================================================
add_heading(doc, "A. Ekler", level=1)

add_heading(doc, "A.1 Teknoloji Yığını", level=2)
add_table(
    doc,
    ["Katman", "Teknoloji", "Sürüm"],
    [
        ["Masaüstü", "Electron + electron-builder", "34.5.8 / 24.13.3"],
        ["Frontend", "React + Vite", "18 / 5.4"],
        ["Routing", "React Router (HashRouter)", "6"],
        ["Stil", "Tailwind benzeri utility CSS", "—"],
        ["Grafikler", "Recharts", "2"],
        ["Backend", "Node.js + Express", "20+ / 4"],
        ["Veritabanı", "better-sqlite3", "11"],
        ["Realtime", "Socket.io", "4"],
        ["Auth", "jsonwebtoken + bcryptjs", "9 / 2"],
        ["Doğrulama", "Zod", "3"],
        ["Loglama", "Pino + pino-http", "9"],
        ["Hata İzleme", "Sentry (node + react)", "8"],
        ["Test", "Vitest + Supertest + Playwright", "1 / 7 / 1"],
        ["Lint", "ESLint flat config", "10"],
        ["CallerID", "C# / .NET 8 self-contained", "—"],
        ["CI", "GitHub Actions", "—"],
    ],
    col_widths=[Cm(3.5), Cm(8.5), Cm(4.5)],
)

add_heading(doc, "A.2 Kurulum Talimatı (Özet)", level=2)
add_bullet(doc, "Yöntem A — GitHub Releases'tan indir: Restoran-POS-v1.1.0-demo.zip → çıkar → Restoran POS.exe çift tıkla.")
add_bullet(doc, "Yöntem B — Geliştirici modu: npm install (kök + server + client) → npm run db:seed → npm run dev → tarayıcı http://localhost:5173.")
add_bullet(doc, "Yöntem C — Yeniden paketleme: npm run dist:prepare → win-unpacked klasörü oluşturulur → manuel ZIP.")

add_heading(doc, "A.3 Demo Kimlik Bilgileri", level=2)
add_table(
    doc,
    ["Rol", "E-posta", "Şifre"],
    [
        ["Yönetici", "admin@demo.com", "123456"],
        ["Kasiyer", "kasiyer@demo.com", "123456"],
        ["Garson", "garson@demo.com", "123456"],
        ["Mutfak", "mutfak@demo.com", "123456"],
    ],
    col_widths=[Cm(4), Cm(7), Cm(5.5)],
)
add_para(
    doc,
    "Not: Bu kimlikler yalnızca demo/seed verisi için geçerlidir. Pilot kurulumda kullanılmaz.",
    italic=True,
)

add_heading(doc, "A.4 Proje Bilgileri", level=2)
add_table(
    doc,
    ["Alan", "Değer"],
    [
        ["Proje Adı", "Restoran POS v1.1.0"],
        ["Geliştirici", "İlhan AVCI (220357008)"],
        ["Üniversite", "Atatürk Üniversitesi"],
        ["Fakülte", "İktisadi ve İdari Bilimler Fakültesi"],
        ["Bölüm", "Yönetim Bilişim Sistemleri"],
        ["Proje Türü", "Bitirme Projesi"],
        ["Versiyon", "1.1.0"],
        ["Test Sayısı", "452 (430 backend + 22 frontend)"],
        ["Lint Uyarısı", "0"],
        ["Migration Sayısı", "13"],
        ["Sürüm Tarihi", "Mayıs 2026"],
        ["Repository", "https://github.com/ilhanavc/restoran-pos"],
        ["Release", "https://github.com/ilhanavc/restoran-pos/releases/tag/v1.1.0-graduation"],
        ["İletişim", "ilhanavci499@gmail.com"],
    ],
    col_widths=[Cm(5), Cm(11.5)],
)

# ---------- Save ----------
output_path = r"D:\dev\restoran-pos-v3\docs\SRS-Restoran-POS-v1.1.0.docx"
doc.save(output_path)
print(f"OK — yazıldı: {output_path}")
