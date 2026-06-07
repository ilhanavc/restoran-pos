# -*- coding: utf-8 -*-
"""
Restoran POS v1.1.0 — Proje Tanıtım Dosyası
Hocaya yönelik kısa, görsel, profesyonel tanıtım.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

SCREENSHOTS = r"D:\dev\restoran-pos-v3\docs\screenshots"

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_image(doc, filename, caption=None, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def add_metric_table(doc, items):
    t = doc.add_table(rows=len(items), cols=2)
    t.style = "Light Shading Accent 1"
    for i, (k, v) in enumerate(items):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t.rows[i].cells[1].text = v

def page_break(doc):
    doc.add_page_break()


# ---------- Doc setup ----------
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================================================
# KAPAK
# ============================================================
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BİTİRME PROJESİ")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RESTORAN POS")
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Modern Restoran Yönetim Sistemi")
run.font.size = Pt(16)
run.italic = True
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sürüm 1.1.0")
run.font.size = Pt(13)

for _ in range(3):
    doc.add_paragraph()

# Highlight metric karts (basit)
t = doc.add_table(rows=2, cols=3)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
metric_data = [
    [("452", "Otomatik Test"), ("0", "Lint Uyarı"), ("13", "Migration")],
    [("12+", "Sprint"), ("452+", "Test Case"), ("9.3/10", "Kalite Skoru")],
]
for r_idx, row in enumerate(metric_data):
    for c_idx, (val, label) in enumerate(row):
        cell = t.rows[r_idx].cells[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = para.add_run(f"{val}\n")
        r1.bold = True
        r1.font.size = Pt(22)
        r1.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        r2 = para.add_run(label)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hazırlayan: İlhan AVCİ")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Öğrenci No: 220357008")
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Atatürk Üniversitesi · İktisadi ve İdari Bilimler Fakültesi")
run.font.size = Pt(11)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Yönetim Bilişim Sistemleri Bölümü")
run.font.size = Pt(11)
run.italic = True

page_break(doc)

# ============================================================
# 1. HOCAYA HİTAP
# ============================================================
add_heading(doc, "Sayın Hocam,", level=1)
add_para(
    doc,
    "Bu doküman; Atatürk Üniversitesi İktisadi ve İdari Bilimler Fakültesi "
    "Yönetim Bilişim Sistemleri Bölümü bitirme projem olan Restoran POS "
    "yazılımının kapsamlı bir tanıtımıdır. Yaklaşık bir yıllık geliştirme "
    "süresince ortaya çıkan bu sistem; gerçek bir restoran işletmesinin "
    "ihtiyaçlarını karşılayacak şekilde tasarlandı ve pilot kullanıma sunuldu.",
)
add_para(
    doc,
    "Bu projede sadece bir akademik gereksinimi karşılamayı değil; aynı zamanda "
    "gerçek dünyada kullanılabilir, kararlı, güvenli ve sürdürülebilir bir "
    "yazılım ürünü ortaya koymayı hedefledim. 12 sprint boyunca uygulanan "
    "modern yazılım mühendisliği prensipleri, kapsamlı test kapsamı (452 "
    "otomatik test), sıfır lint uyarısı ve detaylı dokümantasyon ile bu "
    "hedefin büyük ölçüde başarıldığını umuyorum.",
)
add_para(
    doc,
    "Aşağıdaki bölümlerde projenin ana çıktıları, kullanılan teknolojiler, "
    "uygulanan mühendislik pratikleri ve görsel sunumu yer almaktadır. "
    "Dilerseniz takdim ettiğim ekteki kullanım kılavuzunu ve SRS dokümanını "
    "inceleyerek; ardından paylaştığım GitHub deposundan uygulamayı doğrudan "
    "test edebilirsiniz.",
)
add_para(
    doc,
    "İncelemeleriniz ve değerli geri bildirimleriniz için şimdiden teşekkür ederim.",
)
add_para(doc, "")
add_para(doc, "Saygılarımla,", italic=True)
add_para(doc, "İlhan AVCİ", bold=True)
add_para(doc, "220357008", italic=True)
page_break(doc)

# ============================================================
# 2. PROJE NEDİR?
# ============================================================
add_heading(doc, "1. Proje Nedir?", level=1)
add_para(
    doc,
    "Restoran POS; restoran işletmecilerinin günlük operasyonel ihtiyaçlarını "
    "tek bir uygulama altında karşılayan, masaüstü ortamda çalışan bir Satış "
    "Noktası (Point-of-Sale) sistemidir. Sistem; masa yönetimi, sipariş alma, "
    "ödeme, mutfak yönlendirme, müşteri tanıma, paket sipariş, raporlama ve "
    "yedekleme gibi tüm modülleri içerir.",
)

add_para(doc, "Çözülen Problemler", bold=True, size=14)
add_bullet(doc, "Mevcut ticari POS sistemlerinin (SambaPOS, Adisyo) yüksek aylık abonelik ücretleri")
add_bullet(doc, "Türkçe karakterlerin termal yazıcılarda bozuk çıkması (PC857 ile çözüldü)")
add_bullet(doc, "Telefonla gelen müşterilerin manuel kayıt zorunluluğu (CallerID ile otomatikleştirildi)")
add_bullet(doc, "Tek bilgisayar kurulumlarında veri kaybı riski (otomatik yedekleme + SHA-256)")
add_bullet(doc, "Mutfak ile kasa arasındaki gecikme (Socket.io ile anlık senkronizasyon)")

add_para(doc, "Hedef Kitle", bold=True, size=14)
add_bullet(doc, "Küçük ve orta ölçekli restoranlar, cafeler, fast-food işletmeleri")
add_bullet(doc, "Tek şube veya tek kasa noktası kullanan işletmeler")
add_bullet(doc, "Açık kaynak ve özelleştirilebilir bir çözüm arayanlar")
page_break(doc)

# ============================================================
# 3. ÖZELLIKLER VITRINI (görsel)
# ============================================================
add_heading(doc, "2. Sistem Görselleri", level=1)

add_heading(doc, "2.1 Giriş Ekranı", level=2)
add_para(
    doc,
    "Rol bazlı çoklu kullanıcı sistemi. Demo hızlı giriş kartları ile farklı "
    "rollerden anında giriş yapılabilir. Şifremi unuttum akışı ve zorunlu "
    "şifre değiştirme desteği vardır.",
)
add_image(doc, "01-login.png", caption="Şekil 1: Giriş ekranı, 4 demo hızlı giriş kartı")

page_break(doc)

add_heading(doc, "2.2 Masa Yönetimi", level=2)
add_para(
    doc,
    "Restoranın tüm masaları alan-bazlı sekmelerle görüntülenir. Doluluk durumu "
    "renk kodlarıyla anında ayırt edilir. Masalar arası sipariş transferi ve "
    "kapasite yönetimi desteklenir.",
)
add_image(doc, "02-tables.png", caption="Şekil 2: 4 alan (İç Salon, Bahçe, VIP, Üst Kat) ve 24 masa")
page_break(doc)

add_heading(doc, "2.3 Sipariş Akışı", level=2)
add_para(
    doc,
    "Kategori → ürün → sepet hiyerarşisi ile hızlı sipariş alma. Her ürün için "
    "modifier (acılı, az pişmiş gibi) ve özel not desteği. Mutfağa gönderme "
    "tek tıkla yapılır ve Socket.io ile mutfak ekranı anlık güncellenir.",
)
add_image(doc, "03-order.png", caption="Şekil 3: Sipariş ekranı - kategori, ürün, sepet")
page_break(doc)

add_heading(doc, "2.4 Ödeme Ekranı", level=2)
add_para(
    doc,
    "Nakit/kart/karışık ödeme tipleri; hızlı tutar butonları (50/100/200/500 TL); "
    "indirim uygulama; para üstü otomatik hesabı; çoklu ödeme desteği. "
    "Ödeme sonrası kasa fişi otomatik yazıcıdan basılır.",
)
add_image(doc, "04-payment.png", caption="Şekil 4: Ödeme ekranı")
page_break(doc)

add_heading(doc, "2.5 Mutfak Ekranı", level=2)
add_para(
    doc,
    "Mutfak personeli için optimize edilmiş ekran. Aktif siparişler kart "
    "formatında listelenir. Yaş bazlı renk uyarıları (10 dk sarı, 20 dk "
    "kırmızı) ile aciliyet net olarak gösterilir.",
)
add_image(doc, "05-kitchen.png", caption="Şekil 5: Mutfak ekranı - real-time sipariş takibi")
page_break(doc)

add_heading(doc, "2.6 Paket Sipariş", level=2)
add_para(
    doc,
    "Telefonla gelen siparişler için özel ekran. CallerID entegrasyonu sayesinde "
    "kayıtlı müşteri arıyorsa adresleri otomatik açılır. Yeni müşteri için "
    "hızlı kayıt formu içerir. Ödeme tipi seçimi zorunludur.",
)
add_image(doc, "06-takeaway.png", caption="Şekil 6: Paket sipariş ekranı - müşteri arama")
page_break(doc)

add_heading(doc, "2.7 Müşteri Yönetimi", level=2)
add_para(
    doc,
    "360° müşteri profili: toplam harcama, sipariş sayısı, son ziyaret, "
    "en sevdiği ürünler, sipariş geçmişi, çoklu telefon ve adresler. "
    "Excel/CSV ile toplu import-export.",
)
add_image(doc, "07-customers.png", caption="Şekil 7: Müşteri listesi ekranı")
page_break(doc)

add_heading(doc, "2.8 Raporlar", level=2)
add_para(
    doc,
    "4 etkileşimli grafik: saatlik satış, kategori dağılımı, ödeme tipi, "
    "en çok satan ürünler. X/Z dönem kapatma desteği. Excel ve PDF dışa "
    "aktarma. Kullanıcı ve şube bazlı filtreleme.",
)
add_image(doc, "08-reports.png", caption="Şekil 8: Raporlar ekranı")
page_break(doc)

add_heading(doc, "2.9 Menü Tanımları", level=2)
add_para(
    doc,
    "Sol-sağ panel düzeni ile menü yönetimi. Kategoriler için ikon ve renk "
    "seçimi. Ürünler için resim, modifier grupları, kombo menü desteği. "
    "Akıllı silme: geçmiş kayıt varsa soft-delete, yoksa hard-delete.",
)
add_image(doc, "09-menu-mgmt.png", caption="Şekil 9: Menü tanımları ekranı")
page_break(doc)

add_heading(doc, "2.10 Ayarlar Merkezi", level=2)
add_para(
    doc,
    "Tüm sistem yapılandırması tek bir kart düzeninde toplandı: işletme "
    "bilgileri, kullanıcılar, yazıcılar, yedekleme, denetim, sürüm notları, "
    "Store Bridge durumu, cihaz eşleştirme.",
)
add_image(doc, "10-settings.png", caption="Şekil 10: Ayarlar ana ekranı")
page_break(doc)

add_heading(doc, "2.11 Yedekleme", level=2)
add_para(
    doc,
    "Otomatik (gece 02:00) ve manuel yedek alma. SHA-256 hash ile bütünlük "
    "doğrulama. 30 gün otomatik rotation. Geri yükleme sırasında otomatik "
    "safety revert. Windows Görev Zamanlayıcı ile harici yedek entegrasyonu.",
)
add_image(doc, "11-backup.png", caption="Şekil 11: Yedekleme ekranı")
page_break(doc)

add_heading(doc, "2.12 Denetim Kayıtları (Audit Log)", level=2)
add_para(
    doc,
    "Tüm önemli değişikliklerin (kim, ne, ne zaman) tutulduğu denetim "
    "ekranı. Before/After JSON karşılaştırma görünümü. Filtreleme ve "
    "arama desteği. Mali güvence için kritik özellik.",
)
add_image(doc, "12-audit-log.png", caption="Şekil 12: Denetim kayıtları ekranı")
page_break(doc)

# ============================================================
# 4. TEKNOLOJI YİĞINI
# ============================================================
add_heading(doc, "3. Teknoloji Yığını ve Mimari", level=1)
add_para(
    doc,
    "Sistem modern yazılım geliştirme pratiklerini takip eden 3-katmanlı "
    "(presentation / application / data) bir mimari üzerine kurulmuştur. "
    "Tüm bileşenler aynı Windows makinesinde çalışır, internet bağlantısı "
    "olmadan tam fonksiyonel kullanım sağlar.",
)

add_para(doc, "Katmanlar ve Teknolojiler", bold=True, size=14)
t = doc.add_table(rows=8, cols=2)
t.style = "Light Grid Accent 1"
tech_rows = [
    ["Sunum Katmanı", "React 18 + Vite 5 + HashRouter + Recharts"],
    ["Uygulama Katmanı", "Express 4 + Zod + jsonwebtoken + bcryptjs"],
    ["Gerçek-zamanlı", "Socket.io 4 (kitchen/tables/takeaway kanalları)"],
    ["Veri Katmanı", "better-sqlite3 (WAL kipi, 13 numaralı migration)"],
    ["Donanım Katmanı", "Store Bridge (ESC/POS PC857 encoding)"],
    ["CallerID", "C# / .NET 8 self-contained binary"],
    ["Masaüstü", "Electron 34 + electron-builder 24.13.3"],
    ["Gözlemlenebilirlik", "Pino (NDJSON) + Sentry (redact'li)"],
]
for r_idx, (k, v) in enumerate(tech_rows):
    t.rows[r_idx].cells[0].text = k
    t.rows[r_idx].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[r_idx].cells[1].text = v

add_para(doc, "")
add_para(doc, "Test ve Kalite Stratejisi", bold=True, size=14)
add_bullet(doc, "Vitest + Supertest: 430 backend, 22 frontend birim ve entegrasyon testi")
add_bullet(doc, "Playwright: 2 senaryo E2E testi (table-order-payment + takeaway)")
add_bullet(doc, "GitHub Actions CI: lint + test + e2e + build (concurrency + timeout-minutes)")
add_bullet(doc, "ESLint flat config: --max-warnings 0 (sıfır warning zorunlu)")
add_bullet(doc, "Sentry source map upload (release tag eşleştirmesi)")
page_break(doc)

# ============================================================
# 5. GELIŞTIRME SÜRECI VE METRIKLER
# ============================================================
add_heading(doc, "4. Geliştirme Süreci ve Metrikler", level=1)
add_para(
    doc,
    "Proje, kontrollü 12 sprint + ek sertleştirme fazları (D-1...D-5, DB-1...DB-4, "
    "C-1, C-2, M-1.1, M-1.2, FAZ 0) ile geliştirilmiştir. Her sprint sonunda "
    "kalite metrikleri raporlanmıştır.",
)

add_para(doc, "Genel Metrikler", bold=True, size=14)
add_metric_table(doc, [
    ("Toplam Test Sayısı", "452 (430 backend + 22 frontend)"),
    ("Test Dosya Sayısı", "39 backend + 2 frontend = 41 dosya"),
    ("ESLint Warning", "0 (--max-warnings 0 ile CI gate)"),
    ("Migration Sayısı", "13 (0000-0012, forward-only)"),
    ("Kod Satırı (LOC)", "Yaklaşık 35.000+ (server + client + electron)"),
    ("REST Endpoint Sayısı", "60+"),
    ("Veritabanı Tablo Sayısı", "30+"),
    ("Geliştirme Süresi", "~12 ay (Mayıs 2025 - Mayıs 2026)"),
    ("Runbook Sayısı", "5 operasyonel runbook"),
    ("Genel Kalite Skoru", "9.3 / 10"),
])

add_para(doc, "")
add_para(doc, "Yapılan Sprint Sonu Pasoları", bold=True, size=14)
add_bullet(doc, "Sprint 1-6: Güvenlik, yazıcı, UX, gerçek-zamanlı sistem")
add_bullet(doc, "Sprint 7-8: Test kapsamı (65 → 285 test), production hardening")
add_bullet(doc, "Sprint 9-12: Feature completion, audit, backup-restore hardening")
add_bullet(doc, "D-1...D-5: CI, signing, observability, monolithic decomposition, missing features")
add_bullet(doc, "DB-1...DB-4: Migration disiplini, audit trail, integer minor unit, snapshot")
add_bullet(doc, "C-1, C-2: Railway cloud deployment, refresh token + mobile session")
add_bullet(doc, "M-1.1, M-1.2: Mobile waiter endpoints, device pairing (QR + device_id)")
add_bullet(doc, "FAZ 0 (0.1-0.9): JWT secret guard, TRUST_PROXY, rate-limit, CORS, password policy, Pino, Sentry, secrets audit, CI hardening")

page_break(doc)

# ============================================================
# 6. GÜVENLIK
# ============================================================
add_heading(doc, "5. Güvenlik Önlemleri", level=1)
add_para(
    doc,
    "Sistemde OWASP Top 10 ve KVKK/GDPR uyumlu olacak şekilde, çok katmanlı "
    "savunma yaklaşımıyla güvenlik önlemleri uygulanmıştır.",
)

add_para(doc, "Uygulanan Güvenlik Katmanları", bold=True, size=14)
add_bullet(doc, "🔐 Bcrypt cost 10 ile şifre hash'leme (plain text yok)")
add_bullet(doc, "🔑 JWT (1 saat access) + Refresh Token (30 gün, SHA-256 hash)")
add_bullet(doc, "🚫 Rate limiting (auth: 5/15dk, admin: 20/dk, global: 100/dk)")
add_bullet(doc, "✅ Tüm endpoint Zod ile input validation")
add_bullet(doc, "💉 better-sqlite3 prepared statements (SQL injection kapalı)")
add_bullet(doc, "🌐 CORS production whitelist + dev için LAN regex")
add_bullet(doc, "🛡️ JWT_SECRET min 32 char zorunlu (production fail-fast)")
add_bullet(doc, "📋 Şifre politikası: min 8 + büyük harf + rakam (yeni hesap)")
add_bullet(doc, "🔒 Zorunlu şifre değiştirme (admin reset sonrası)")
add_bullet(doc, "📝 Audit trail (entity_mutations) - before/after JSON")
add_bullet(doc, "🎭 Sentry redact (password, token, cookie, jwtSecret, bridgeToken)")
add_bullet(doc, "📹 Session replay: maskAllText + maskAllInputs + blockAllMedia")
add_bullet(doc, "🆔 X-Request-Id UUID format doğrulaması (header injection guard)")
add_bullet(doc, "⏱️ Print queue lease-based claim (ownership guard)")
add_bullet(doc, "🚦 trust proxy hops env-driven (geçersizde fallback 0)")
add_bullet(doc, "🔍 Audit log viewer ile değişiklik izleme")
page_break(doc)

# ============================================================
# 7. BU PROJEYI ÖZEL KILAN NEDIR?
# ============================================================
add_heading(doc, "6. Projenin Akademik ve Pratik Değeri", level=1)

add_para(doc, "Akademik Açıdan", bold=True, size=14)
add_bullet(doc, "IEEE 830-1998 standardında hazırlanmış kapsamlı SRS dokümanı")
add_bullet(doc, "Yazılım mühendisliği prensipleri (modülerlik, test güdümlü geliştirme)")
add_bullet(doc, "Veritabanı tasarımı ve normalleştirme (3NF, FK ilişkileri)")
add_bullet(doc, "Çok katmanlı mimari tasarımı (3-tier)")
add_bullet(doc, "Tasarım desenleri: Service Layer, Repository, Strategy, Observer (Socket.io)")
add_bullet(doc, "DevOps pratikleri: CI/CD, lint gate, automated testing")
add_bullet(doc, "Güvenlik mühendisliği (OWASP, KVKK uyum)")
add_bullet(doc, "Operasyonel gözlemlenebilirlik (logging, monitoring, error tracking)")

add_para(doc, "Pratik Açıdan", bold=True, size=14)
add_bullet(doc, "Gerçek bir restoranda kullanılmaya hazır (pilot kullanım)")
add_bullet(doc, "Açık kaynak alternatifi olarak ticari rakiplere meydan okuyabilir")
add_bullet(doc, "Türkçe karakter sorununu çözen ESC/POS PC857 entegrasyonu")
add_bullet(doc, "Çevrimdışı çalışma garantisi (internet kopuşunda iş kesintisiz devam eder)")
add_bullet(doc, "Otomatik yedekleme ve geri yükleme ile veri güvenliği")

add_para(doc, "Kalite Göstergeleri", bold=True, size=14)
add_bullet(doc, "452 otomatik test (>%95 kritik akış kapsamı)")
add_bullet(doc, "Sıfır lint warning (kod kalitesi standardı)")
add_bullet(doc, "13 numaralı migration (veri evrimi kontrolü)")
add_bullet(doc, "5 operasyonel runbook (saha kullanımı için)")
add_bullet(doc, "9.3/10 genel kalite skoru")
page_break(doc)

# ============================================================
# 8. ERIŞIM VE TESLIM
# ============================================================
add_heading(doc, "7. Projeye Erişim", level=1)

add_para(doc, "GitHub Repository", bold=True, size=14)
add_para(doc, "Tüm kaynak kod, dokümanlar ve tarihçe:")
add_para(doc, "🔗 https://github.com/ilhanavc/restoran-pos", color=RGBColor(0x4F, 0x46, 0xE5))

add_para(doc, "")
add_para(doc, "Bitirme Teslim Sürümü (Hazır Paket)", bold=True, size=14)
add_para(doc, "Aşağıdaki linkten doğrudan ZIP indirip Restoran POS.exe ile çalıştırabilirsiniz:")
add_para(doc, "🔗 https://github.com/ilhanavc/restoran-pos/releases/tag/v1.1.0-graduation",
         color=RGBColor(0x4F, 0x46, 0xE5))

add_para(doc, "")
add_para(doc, "Bu Teslim Paketinde Bulunan Dosyalar", bold=True, size=14)
add_bullet(doc, "📄 PROJE-TANITIM.docx (bu doküman)")
add_bullet(doc, "📄 SRS-Restoran-POS-v1.1.0.docx (Yazılım Gereksinimleri Belirtimi)")
add_bullet(doc, "📄 KULLANIM-KILAVUZU-Restoran-POS.docx (Adım adım kullanıcı kılavuzu)")
add_bullet(doc, "📦 Restoran-POS-v1.1.0-demo.zip (Çalıştırılabilir uygulama + demo verisi)")
add_bullet(doc, "📸 screenshots/ (12 ekran görüntüsü)")
add_bullet(doc, "📁 docs/ (Tüm proje dokümantasyonu)")

add_para(doc, "")
add_para(doc, "Demo Kimlik Bilgileri", bold=True, size=14)
t = doc.add_table(rows=5, cols=3)
t.style = "Light Grid Accent 1"
hdr = ["Rol", "E-posta", "Şifre"]
demo_rows = [
    ["Yönetici", "admin@demo.com", "123456"],
    ["Kasiyer", "kasiyer@demo.com", "123456"],
    ["Garson", "garson@demo.com", "123456"],
    ["Mutfak", "mutfak@demo.com", "123456"],
]
for i, h in enumerate(hdr):
    cell = t.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for r_idx, row in enumerate(demo_rows, start=1):
    for c_idx, val in enumerate(row):
        t.rows[r_idx].cells[c_idx].text = val
page_break(doc)

# ============================================================
# 9. KAPANIŞ
# ============================================================
add_heading(doc, "8. Sonuç ve Değerlendirme", level=1)
add_para(
    doc,
    "Restoran POS, akademik bir bitirme projesi olmanın yanında, gerçek bir "
    "işletmede kullanılabilecek olgunluğa erişmiş bir yazılım ürünüdür. "
    "12 sprint boyunca uygulanan yazılım mühendisliği prensipleri, kapsamlı "
    "test stratejisi ve titiz dokümantasyon ile profesyonel bir geliştirme "
    "deneyimi sergilemektedir.",
)

add_para(
    doc,
    "Projenin gelecek yol haritasında çoklu şube desteği, mobil garson uygulaması, "
    "ödeme terminal SDK entegrasyonu ve resmi e-belge yetkilendirmesi yer "
    "almaktadır. Ancak bu çıkarımlar, bitirme teslim sürümünün kapsamı dışında "
    "tutulmuş, ileri sürümler için yol haritasında belirtilmiştir.",
)

add_para(
    doc,
    "Eleştiri, soru veya öneri için her zaman ulaşabileceğiniz iletişim "
    "kanallarımı aşağıda bulabilirsiniz:",
)

add_para(doc, "")
add_para(doc, "İletişim Bilgileri", bold=True, size=14)
add_bullet(doc, "📧 E-posta: ilhanavci499@gmail.com")
add_bullet(doc, "🐙 GitHub: github.com/ilhanavc")
add_bullet(doc, "📚 Proje Repo: github.com/ilhanavc/restoran-pos")
add_bullet(doc, "🎓 Üniversite: Atatürk Üniversitesi · İİBF · Yönetim Bilişim Sistemleri")
add_bullet(doc, "🆔 Öğrenci No: 220357008")

add_para(doc, "")
add_para(doc, "")
add_para(
    doc,
    "Değerli zamanınızı ayırarak bu çalışmayı incelediğiniz için teşekkür ederim.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=12,
)
add_para(doc, "")
add_para(
    doc,
    "İlhan AVCİ",
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=14,
)
add_para(
    doc,
    "Mayıs 2026",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

# Save
output_path = r"D:\dev\restoran-pos-v3\docs\PROJE-TANITIM-Restoran-POS.docx"
doc.save(output_path)
print(f"OK — yazıldı: {output_path}")
