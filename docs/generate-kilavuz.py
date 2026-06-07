# -*- coding: utf-8 -*-
"""
Restoran POS v1.1.0 — Kullanım Kılavuzu (DOCX)
Ekran görüntüleri ile adım adım kullanıcı dokümanı.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

SCREENSHOTS = r"D:\dev\restoran-pos-v3\docs\screenshots"

# ---------- Helpers ----------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if align is not None:
        p.alignment = align
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_image(doc, filename, caption=None, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Şekil: {caption}")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p

def add_info_box(doc, title, text, color="info"):
    """Bilgi kutusu (tablo ile)"""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Light Shading Accent 1"
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(f"💡 {title}\n")
    run.bold = True
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def page_break(doc):
    doc.add_page_break()


# ---------- Doc setup ----------
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================================================
# KAPAK
# ============================================================
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RESTORAN POS")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Kullanım Kılavuzu")
run.bold = True
run.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("v1.1.0")
run.font.size = Pt(14)
run.italic = True

for _ in range(2):
    doc.add_paragraph()

# Logo benzeri kart
add_para(
    doc,
    "Modern Restoran Yönetim ve Satış Noktası Sistemi",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    italic=True,
    size=13,
)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hazırlayan: İlhan AVCI")
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("220357008")
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Atatürk Üniversitesi · İİBF · Yönetim Bilişim Sistemleri")
run.font.size = Pt(11)
run.italic = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mayıs 2026")
run.font.size = Pt(11)

page_break(doc)

# ============================================================
# İÇİNDEKİLER
# ============================================================
add_heading(doc, "İçindekiler", level=1)
toc_items = [
    "1. Sistem Hakkında",
    "2. Kurulum ve İlk Açılış",
    "3. Sisteme Giriş Yapma",
    "4. Masa Yönetimi",
    "5. Sipariş Alma",
    "6. Ödeme İşlemi",
    "7. Mutfak Ekranı",
    "8. Paket Sipariş",
    "9. Müşteri Yönetimi",
    "10. Raporlar ve Analizler",
    "11. Menü Tanımları",
    "12. Ayarlar",
    "13. Yedekleme ve Geri Yükleme",
    "14. Denetim Kayıtları (Audit Log)",
    "15. Sık Karşılaşılan Sorular",
    "16. Destek ve İletişim",
]
for label in toc_items:
    add_bullet(doc, label)

page_break(doc)

# ============================================================
# 1. SİSTEM HAKKINDA
# ============================================================
add_heading(doc, "1. Sistem Hakkında", level=1)
add_para(
    doc,
    "Restoran POS, modern restoran işletmeleri için tasarlanmış, masaüstü ortamda "
    "çalışan tam özellikli bir Satış Noktası (Point-of-Sale) sistemidir. "
    "Sistem; masa yönetimi, sipariş alma, ödeme işlemleri, mutfak yönlendirme, "
    "müşteri takibi, paket sipariş ve detaylı raporlama gibi tüm restoran operasyonlarını "
    "tek bir uygulama altında toplar.",
)

add_para(doc, "Sistem Özellikleri", bold=True, size=13)
add_bullet(doc, "🪑 Görsel masa düzeni ve doluluk takibi")
add_bullet(doc, "📋 Hızlı sipariş alma akışı")
add_bullet(doc, "💳 Nakit/kart/karışık ödeme ve hızlı tutar butonları")
add_bullet(doc, "👨‍🍳 Anlık (real-time) mutfak ekranı senkronizasyonu")
add_bullet(doc, "🥡 Paket sipariş ve müşteri tanıma (CallerID)")
add_bullet(doc, "🧾 Türkçe karakter destekli ESC/POS termal yazıcı çıktısı")
add_bullet(doc, "📊 İnteraktif raporlar ve grafikler")
add_bullet(doc, "💾 Otomatik veritabanı yedekleme")
add_bullet(doc, "🔐 Çoklu rol bazlı yetkilendirme (Yönetici, Kasiyer, Garson, Mutfak)")

add_para(doc, "Sistem Gereksinimleri", bold=True, size=13)
add_bullet(doc, "İşletim Sistemi: Windows 10 (1809+) veya Windows 11")
add_bullet(doc, "İşlemci: 64-bit Intel/AMD (Core i3 veya üzeri önerilir)")
add_bullet(doc, "Bellek: En az 4 GB RAM")
add_bullet(doc, "Disk: En az 2 GB boş alan")
add_bullet(doc, "Ek Donanım: ESC/POS termal yazıcı (opsiyonel), CallerID HID cihazı (opsiyonel)")
page_break(doc)

# ============================================================
# 2. KURULUM VE İLK AÇILIŞ
# ============================================================
add_heading(doc, "2. Kurulum ve İlk Açılış", level=1)
add_para(doc, "2.1 Uygulamayı İndirme", bold=True, size=13)
add_number(doc, "GitHub Releases sayfasından Restoran-POS-v1.1.0-demo.zip dosyasını indirin: github.com/ilhanavc/restoran-pos/releases")
add_number(doc, "İndirilen ZIP dosyasını sağ tıklayıp 'Tümünü Çıkar' seçeneği ile bir klasöre çıkarın.")
add_number(doc, "Çıkarılan klasör içerisindeki Restoran POS.exe dosyasına çift tıklayın.")

add_info_box(
    doc,
    "Önemli Not",
    "İlk açılışta Windows SmartScreen güvenlik uyarısı gösterebilir. "
    "Bu, uygulamanın imzalanmamış olmasından kaynaklanır (akademik proje). "
    "'Diğer bilgiler' → 'Yine de çalıştır' seçeneği ile devam edebilirsiniz.",
)

add_para(doc, "2.2 İlk Açılış", bold=True, size=13)
add_para(
    doc,
    "Uygulama ilk açıldığında, demo veri tabanı otomatik olarak yüklenir. "
    "Bu sayede demo kullanıcı bilgileriyle sisteme hemen giriş yapabilirsiniz.",
)
page_break(doc)

# ============================================================
# 3. SİSTEME GİRİŞ
# ============================================================
add_heading(doc, "3. Sisteme Giriş Yapma", level=1)
add_para(
    doc,
    "Uygulama açıldığında giriş ekranı sizi karşılar. Buradan e-posta adresiniz ve "
    "şifrenizle sisteme erişebilirsiniz.",
)
add_image(doc, "01-login.png", caption="Giriş Ekranı")

add_para(doc, "3.1 Demo Kullanıcılar", bold=True, size=13)
add_para(doc, "Sistem ilk açıldığında dört farklı demo kullanıcı kullanıma hazırdır:")

t = doc.add_table(rows=5, cols=3)
t.style = "Light Grid Accent 1"
header = ["Rol", "E-posta", "Şifre"]
rows = [
    ["Yönetici", "admin@demo.com", "123456"],
    ["Kasiyer", "kasiyer@demo.com", "123456"],
    ["Garson", "garson@demo.com", "123456"],
    ["Mutfak", "mutfak@demo.com", "123456"],
]
for i, h in enumerate(header):
    cell = t.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
for r_idx, row in enumerate(rows, start=1):
    for c_idx, val in enumerate(row):
        t.rows[r_idx].cells[c_idx].text = val

add_para(doc, "3.2 Adım Adım Giriş", bold=True, size=13)
add_number(doc, "E-posta alanına kullanıcı e-postanızı yazın (örn: admin@demo.com)")
add_number(doc, "Şifre alanına şifrenizi yazın (örn: 123456)")
add_number(doc, "'Giriş Yap' butonuna tıklayın")
add_number(doc, "Veya altta bulunan 'DEMO HIZLI GİRİŞ' kartlarından birine tıklayarak hızlıca giriş yapın")

add_info_box(
    doc,
    "Şifremi Unuttum",
    "Şifrenizi unutursanız 'Şifremi unuttum' bağlantısına tıklayarak sıfırlama "
    "talebi gönderebilirsiniz. Yönetici onayından sonra yeni şifre belirleyebilirsiniz.",
)
page_break(doc)

# ============================================================
# 4. MASA YÖNETİMİ
# ============================================================
add_heading(doc, "4. Masa Yönetimi", level=1)
add_para(
    doc,
    "Giriş yaptıktan sonra ilk karşınıza çıkan ekran masa yönetimi ekranıdır. "
    "Burada restoranınızın tüm masalarını tek bakışta görebilir, doluluk durumlarını "
    "renkler ile ayırt edebilirsiniz.",
)
add_image(doc, "02-tables.png", caption="Masa Yönetimi Ekranı")

add_para(doc, "4.1 Alan ve Masa Düzeni", bold=True, size=13)
add_para(doc, "Ekran üst kısımda restoranın farklı bölgeleri sekmeler halinde listelenir:")
add_bullet(doc, "İç Salon — 10 masa")
add_bullet(doc, "Bahçe — 6 masa")
add_bullet(doc, "VIP — 3 masa")
add_bullet(doc, "Üst Kat — 5 masa")

add_para(doc, "4.2 Masa Renk Kodları", bold=True, size=13)
add_bullet(doc, "🟢 Yeşil: Boş masa (sipariş için hazır)")
add_bullet(doc, "🟡 Sarı: Sipariş var, hazırlanıyor")
add_bullet(doc, "🔴 Kırmızı: Uzun süredir aktif (uyarı durumu)")
add_bullet(doc, "🔵 Mavi: Ödemeye geçildi")

add_para(doc, "4.3 Masa İşlemleri", bold=True, size=13)
add_number(doc, "Bir masaya tıklayarak sipariş ekranını açın")
add_number(doc, "Masayı uzun süre dolu ise masa transferi yaparak başka masaya taşıyın")
add_number(doc, "Üst sağ köşeden 'Masa Bilgisi' ile masanın kapasitesini ve durumunu görün")
page_break(doc)

# ============================================================
# 5. SİPARİŞ ALMA
# ============================================================
add_heading(doc, "5. Sipariş Alma", level=1)
add_para(
    doc,
    "Bir masaya tıkladığınızda sipariş ekranı açılır. Bu ekran, "
    "kategori-ürün hiyerarşisi, sepet ve sipariş kontrol butonlarından oluşur.",
)
add_image(doc, "03-order.png", caption="Sipariş Ekranı")

add_para(doc, "5.1 Ekran Yapısı", bold=True, size=13)
add_bullet(doc, "Sol Panel: Kategori listesi (Çorbalar, Pideler, Ana Yemekler, İçecekler)")
add_bullet(doc, "Orta Panel: Seçilen kategorideki ürünler")
add_bullet(doc, "Sağ Panel: Sepet (eklenmiş ürünler, toplam tutar)")

add_para(doc, "5.2 Adım Adım Sipariş Alma", bold=True, size=13)
add_number(doc, "Sol panelden bir kategori seçin (örn: Pideler)")
add_number(doc, "Orta panelde görünen ürünlerden birine tıklayarak sepete ekleyin")
add_number(doc, "Aynı ürünü birden fazla eklemek için tekrar tıklayın")
add_number(doc, "Sağ panelde sepet otomatik güncellenir, alt kısımda toplam tutar görünür")
add_number(doc, "İhtiyaç halinde 'Modifier' (acılı, az pişmiş gibi) ekleyebilirsiniz")
add_number(doc, "Her ürün için özel not (örn: 'soğansız') girebilirsiniz")
add_number(doc, "'Mutfağa Gönder' butonuna tıklayarak siparişi mutfağa iletin")

add_info_box(
    doc,
    "Gerçek Zamanlı Güncelleme",
    "'Mutfağa Gönder' butonuna basıldığı anda, mutfak ekranı otomatik olarak "
    "güncellenir. Yenileme yapmanıza gerek yoktur. Socket.io teknolojisi sayesinde "
    "tüm değişiklikler anlık olarak ekranlara yansır.",
)

add_para(doc, "5.3 Sipariş Düzenleme", bold=True, size=13)
add_bullet(doc, "Sepetteki bir ürünün üzerine tıklayarak adetini artırabilir/azaltabilirsiniz")
add_bullet(doc, "Üç-nokta menüden ürünü silebilir veya not ekleyebilirsiniz")
add_bullet(doc, "İptal işlemi için 'Sipariş İptal' butonunu kullanın (yetki kontrolü vardır)")
page_break(doc)

# ============================================================
# 6. ÖDEME
# ============================================================
add_heading(doc, "6. Ödeme İşlemi", level=1)
add_para(
    doc,
    "Sipariş tamamlandığında ödeme almak için 'Ödeme' butonuna tıklayın. "
    "Ödeme ekranı, hızlı tutar butonları, ödeme tipi seçimi ve para üstü "
    "hesaplaması ile donatılmıştır.",
)
add_image(doc, "04-payment.png", caption="Ödeme Ekranı")

add_para(doc, "6.1 Ödeme Tipleri", bold=True, size=13)
add_bullet(doc, "Nakit — fiziksel para")
add_bullet(doc, "Kart — kredi/banka kartı (POS cihazı ayrı)")
add_bullet(doc, "Karışık — bir kısmı nakit, bir kısmı kart")

add_para(doc, "6.2 Hızlı Tutar Butonları", bold=True, size=13)
add_para(doc, "50 TL, 100 TL, 200 TL, 500 TL gibi sık kullanılan tutarları tek tıklamayla eklenir.")

add_para(doc, "6.3 Adım Adım Ödeme", bold=True, size=13)
add_number(doc, "Sipariş üzerinden 'Ödeme' butonuna tıklayın")
add_number(doc, "Ödeme ekranında toplam tutar görüntülenir")
add_number(doc, "Müşterinin ödeme tipini seçin (Nakit/Kart)")
add_number(doc, "Müşterinin verdiği tutarı hızlı butonlarla veya manuel girin")
add_number(doc, "Para üstü otomatik hesaplanır")
add_number(doc, "İhtiyaç halinde indirim uygulayın")
add_number(doc, "'Ödemeyi Tamamla' butonuna tıklayın")
add_number(doc, "Masa otomatik olarak boşa düşer, kasa fişi yazıcıdan basılır")
page_break(doc)

# ============================================================
# 7. MUTFAK EKRANI
# ============================================================
add_heading(doc, "7. Mutfak Ekranı", level=1)
add_para(
    doc,
    "Mutfak personeli için optimize edilmiş ekran. Aktif tüm siparişleri "
    "kart formatında gösterir ve sipariş yaşına göre renk uyarıları sağlar.",
)
add_image(doc, "05-kitchen.png", caption="Mutfak Ekranı")

add_para(doc, "7.1 Sipariş Kartı Bileşenleri", bold=True, size=13)
add_bullet(doc, "Masa numarası veya 'PAKET' ibaresi (üst kısım)")
add_bullet(doc, "Sipariş tarihi ve saati")
add_bullet(doc, "Hazırlama süresi (geçen dakika)")
add_bullet(doc, "Ürün listesi (adet ve modifier'lar ile)")
add_bullet(doc, "Hazırlandı / Tamamlandı butonları")

add_para(doc, "7.2 Renkli Yaş Uyarıları", bold=True, size=13)
add_bullet(doc, "⚪ Beyaz/normal: Yeni sipariş (0-10 dakika)")
add_bullet(doc, "🟡 Sarı: Orta gecikme (10-20 dakika)")
add_bullet(doc, "🔴 Kırmızı: Acil! Uzun bekleme (20+ dakika)")

add_para(doc, "7.3 Kalem Bazlı Hazırlık", bold=True, size=13)
add_para(
    doc,
    "Her ürünün yanında 'Hazırlandı' onay kutusu vardır. Bir kalemin hazır "
    "olduğunu işaretlediğinizde, durum anında kasiyer ekranına yansır.",
)

add_info_box(
    doc,
    "Sesli Uyarılar",
    "Yeni bir sipariş geldiğinde mutfak ekranı kısa bir bildirim sesi çalar. "
    "Bu özellik ayarlardan açılıp kapatılabilir.",
)
page_break(doc)

# ============================================================
# 8. PAKET SİPARİŞ
# ============================================================
add_heading(doc, "8. Paket Sipariş", level=1)
add_para(
    doc,
    "Telefonla gelen paket siparişler için özel olarak tasarlanmış ekran. "
    "Müşteri arama, adres seçimi ve ödeme tipi belirleme akışlarını içerir.",
)
add_image(doc, "06-takeaway.png", caption="Paket Sipariş Ekranı")

add_para(doc, "8.1 CallerID Entegrasyonu", bold=True, size=13)
add_para(
    doc,
    "Müşteri telefonu çaldığında, CallerID cihazı numarayı otomatik olarak "
    "sisteme iletir. Kayıtlı bir müşteri arıyorsa adı ve adres bilgileri "
    "otomatik açılır.",
)

add_para(doc, "8.2 Manuel Paket Sipariş", bold=True, size=13)
add_number(doc, "Ana menüden 'Paket' sekmesine geçin")
add_number(doc, "Telefon numarası alanına müşterinin numarasını girin")
add_number(doc, "Sistem otomatik olarak kayıtlı müşteriyi bulur ve adresleri listeler")
add_number(doc, "Yeni müşteri ise 'Yeni Müşteri Ekle' formu açılır")
add_number(doc, "Müşterinin teslimat adresini seçin")
add_number(doc, "Ödeme tipini seçin (zorunlu - Nakit veya Kart)")
add_number(doc, "Ürünleri sepete ekleyin")
add_number(doc, "'Sipariş Oluştur' ile siparişi tamamlayın")

add_info_box(
    doc,
    "Çoklu Adres",
    "Bir müşterinin birden fazla teslimat adresi olabilir (Ev, İş gibi). "
    "Sistem tüm adresleri listelenir ve siparişi yaparken seçilir.",
)
page_break(doc)

# ============================================================
# 9. MÜŞTERİ YÖNETİMİ
# ============================================================
add_heading(doc, "9. Müşteri Yönetimi", level=1)
add_para(
    doc,
    "Müşteri bilgilerini, sipariş geçmişini ve adresleri yönettiğiniz ekran. "
    "Excel/CSV ile toplu import-export desteği bulunur.",
)
add_image(doc, "07-customers.png", caption="Müşteriler Ekranı")

add_para(doc, "9.1 Müşteri Listesi Özellikleri", bold=True, size=13)
add_bullet(doc, "Sayfalama (50 kayıt/sayfa, 'Daha Fazla Yükle' ile genişler)")
add_bullet(doc, "Arama: ad, soyad, telefon")
add_bullet(doc, "Filtreler: toplam sipariş sayısı, son ziyaret tarihi")
add_bullet(doc, "Her satırda hızlı eylem butonları")

add_para(doc, "9.2 Müşteri 360 Profili", bold=True, size=13)
add_para(doc, "Bir müşteriye tıkladığınızda, detaylı profil sayfası açılır:")
add_bullet(doc, "Toplam harcama tutarı")
add_bullet(doc, "Toplam sipariş sayısı")
add_bullet(doc, "Son ziyaret tarihi")
add_bullet(doc, "En çok sipariş ettiği 3 ürün")
add_bullet(doc, "Sipariş geçmişi (tarih + tutar)")
add_bullet(doc, "Kayıtlı telefonlar ve adresler")

add_para(doc, "9.3 Toplu Import / Export", bold=True, size=13)
add_number(doc, "Excel/CSV ile mevcut müşteri verisini içe aktarmak için 'İçe Aktar' butonunu kullanın")
add_number(doc, "Mevcut müşteri listesini Excel olarak dışa aktarmak için 'Dışa Aktar' butonuna tıklayın")
add_number(doc, "Şablon dosya 'Şablon İndir' ile elde edilebilir")
page_break(doc)

# ============================================================
# 10. RAPORLAR
# ============================================================
add_heading(doc, "10. Raporlar ve Analizler", level=1)
add_para(
    doc,
    "İşletmenizin satış performansını ölçtüğünüz, 4 etkileşimli grafik ile "
    "donatılmış rapor ekranı. Excel ve PDF olarak dışa aktarılabilir.",
)
add_image(doc, "08-reports.png", caption="Raporlar Ekranı")

add_para(doc, "10.1 Filtreleme", bold=True, size=13)
add_bullet(doc, "Tarih aralığı: Bugün, Bu Hafta, Bu Ay, Özel Tarih")
add_bullet(doc, "Şube/alan filtresi (varsa)")
add_bullet(doc, "Kullanıcı bazlı filtre")

add_para(doc, "10.2 Görüntülenen Grafikler", bold=True, size=13)
add_bullet(doc, "📈 Saatlik Satış (line chart) — günün hangi saatlerinde yoğunluk var")
add_bullet(doc, "🥧 Kategori Dağılımı (pie chart) — hangi kategoriler en çok satıyor")
add_bullet(doc, "🍩 Ödeme Tipi (donut chart) — nakit/kart oranı")
add_bullet(doc, "📊 En Çok Satan Ürünler (bar chart) — top 10 ürün")

add_para(doc, "10.3 X / Z Rapor", bold=True, size=13)
add_para(
    doc,
    "Klasik POS sistemlerindeki X (ara) ve Z (kapanış) raporu özelliği vardır. "
    "Gün sonu kapanışı için 'Z Raporu Al' butonu kullanılır. Kapatılmış bir gün "
    "için iade işlemi yapılamaz (mali güvence).",
)

add_para(doc, "10.4 Dışa Aktarma", bold=True, size=13)
add_bullet(doc, "Excel formatında detaylı tablolar")
add_bullet(doc, "PDF ile yazdırılabilir özet")
add_bullet(doc, "Grafiklerin görsel kaydı")
page_break(doc)

# ============================================================
# 11. MENÜ TANIMLARI
# ============================================================
add_heading(doc, "11. Menü Tanımları", level=1)
add_para(
    doc,
    "Restoranınızın menüsünü yönetmek için kullanılır. Kategoriler, ürünler, "
    "fiyatlar, modifier'lar ve yazıcı yönlendirmesi burada tanımlanır.",
)
add_image(doc, "09-menu-mgmt.png", caption="Menü Tanımları Ekranı")

add_para(doc, "11.1 Ekran Yapısı", bold=True, size=13)
add_bullet(doc, "Sol Panel: Kategori listesi + 'Ekle' butonu")
add_bullet(doc, "Sağ Panel: Seçilen kategorinin ürünleri")
add_bullet(doc, "Sağ Üst: Ürün arama + 'Yeni Ürün Ekle' butonu")

add_para(doc, "11.2 Kategori Ekleme", bold=True, size=13)
add_number(doc, "Sol panelde '+ Ekle' butonuna tıklayın")
add_number(doc, "Kategori adını girin (örn: 'Tatlılar')")
add_number(doc, "Yazıcı hedefini seçin (Mutfak/Bar/Kasa)")
add_number(doc, "İkonu seçin (15+ seçenek)")
add_number(doc, "Rengi seçin (13 farklı renk)")
add_number(doc, "'Kaydet' ile tamamlayın")

add_para(doc, "11.3 Ürün Ekleme", bold=True, size=13)
add_number(doc, "Bir kategori seçin")
add_number(doc, "Sağ üstteki 'Yeni Ürün Ekle' butonuna tıklayın")
add_number(doc, "Ürün adı, açıklama, fiyat ve KDV oranını girin")
add_number(doc, "Ürün resmi yükleyin (opsiyonel)")
add_number(doc, "Modifier grupları varsa ekleyin (örn: pişme derecesi)")
add_number(doc, "Kombo menü ise ilgili ürünleri seçin")
add_number(doc, "'Kaydet' ile tamamlayın")

add_info_box(
    doc,
    "Akıllı Silme (Hybrid Delete)",
    "Bir kategori veya ürünü silmek istediğinizde sistem otomatik karar verir: "
    "Eğer geçmiş sipariş kaydı yoksa kalıcı silinir. Eğer geçmiş kaydı varsa "
    "soft-delete yapılır (görünmez ama veri kaybı olmaz, raporlar etkilenmez).",
)
page_break(doc)

# ============================================================
# 12. AYARLAR
# ============================================================
add_heading(doc, "12. Ayarlar", level=1)
add_para(
    doc,
    "Sistemin tüm yapılandırmasını yönetebileceğiniz merkez. İşletme bilgileri, "
    "kullanıcı yönetimi, yazıcı tanımları, yedekleme ve daha fazlası burada.",
)
add_image(doc, "10-settings.png", caption="Ayarlar Ana Ekranı")

add_para(doc, "12.1 Ayar Kartları", bold=True, size=13)
add_bullet(doc, "İşletme Bilgileri — ad, telefon, adres, vergi no")
add_bullet(doc, "Kullanıcılar — yeni kullanıcı ekleme, rol değiştirme, şifre sıfırlama")
add_bullet(doc, "Yazıcılar — yazıcı tanımları, IP/USB yapılandırma, test çıktısı")
add_bullet(doc, "Bakım ve Yedekleme — manuel/otomatik yedek, geri yükleme")
add_bullet(doc, "Audit Log — sistem değişiklik kayıtları")
add_bullet(doc, "Sürüm Notları — uygulama versiyonu, güncelleme arama")
add_bullet(doc, "Store Bridge — yazıcı bağlantı durumu, log görüntüleme")
add_bullet(doc, "Cihazlar — mobil/tablet eşleştirme (QR kod ile)")
page_break(doc)

# ============================================================
# 13. YEDEKLEME
# ============================================================
add_heading(doc, "13. Yedekleme ve Geri Yükleme", level=1)
add_para(
    doc,
    "Veri güvenliği için hem otomatik hem manuel yedekleme imkanı sunan ekran. "
    "Yedekler SHA-256 hash ile imzalanır ve bütünlüğü doğrulanır.",
)
add_image(doc, "11-backup.png", caption="Yedekleme Ekranı")

add_para(doc, "13.1 Otomatik Yedek", bold=True, size=13)
add_bullet(doc, "Her gece 02:00'de otomatik tetiklenir")
add_bullet(doc, "Son 30 günün yedekleri saklanır (otomatik rotasyon)")
add_bullet(doc, "İçerik: veritabanı + uploads + yapılandırma + meta.json")

add_para(doc, "13.2 Manuel Yedek Alma", bold=True, size=13)
add_number(doc, "'Manuel Yedek Al' butonuna tıklayın")
add_number(doc, "Yedek hazırlanır (genelde 5-10 saniye)")
add_number(doc, "Başarı bildirimi gösterilir")
add_number(doc, "Liste anlık güncellenir, son yedek en üstte görünür")

add_para(doc, "13.3 Geri Yükleme", bold=True, size=13)
add_number(doc, "Yedekler listesinden bir yedek seçin")
add_number(doc, "'Geri Yükle' butonuna tıklayın")
add_number(doc, "Yedek özeti görüntülenir: tarih, boyut, satır sayıları")
add_number(doc, "Açık sipariş varsa sistem uyarı verir, devam etmemek önerilir")
add_number(doc, "'Onayla ve Geri Yükle' ile işlem başlatılır")
add_number(doc, "SHA-256 doğrulaması ve bütünlük kontrolü yapılır")
add_number(doc, "Başarısız olursa otomatik safety revert ile eski veri geri yüklenir")

add_para(doc, "13.4 Harici Yedek (Windows Görev Zamanlayıcı)", bold=True, size=13)
add_para(
    doc,
    "Yedeklerinizi USB veya ağ sürücüsüne otomatik kopyalamak için Windows "
    "Görev Zamanlayıcı entegrasyonu vardır. Gece 03:00'de robocopy ile harici "
    "hedef klasöre yedek kopyalanır.",
)
page_break(doc)

# ============================================================
# 14. AUDIT LOG
# ============================================================
add_heading(doc, "14. Denetim Kayıtları (Audit Log)", level=1)
add_para(
    doc,
    "Sistemde yapılan tüm önemli değişikliklerin (kim, ne, ne zaman) tutulduğu "
    "denetim ekranı. Mali ve operasyonel sorumluluğun izlenmesi için kritiktir.",
)
add_image(doc, "12-audit-log.png", caption="Denetim Kayıtları")

add_para(doc, "14.1 Kayıt Türleri", bold=True, size=13)
add_bullet(doc, "Ürün değişiklikleri (ekleme, güncelleme, silme)")
add_bullet(doc, "Kategori değişiklikleri")
add_bullet(doc, "Stok hareketleri")
add_bullet(doc, "Müşteri kayıt değişiklikleri")
add_bullet(doc, "İşletme/yazıcı/kullanıcı yapılandırma değişiklikleri")
add_bullet(doc, "Sipariş oluşturma/iptal")
add_bullet(doc, "Ödeme ve iade kayıtları")

add_para(doc, "14.2 Filtreleme", bold=True, size=13)
add_bullet(doc, "Tablo bazlı filtre (sipariş, ödeme, ürün, müşteri)")
add_bullet(doc, "İşlem bazlı filtre (create, update, delete)")
add_bullet(doc, "Tarih aralığı")
add_bullet(doc, "Kullanıcı bazlı filtre")

add_para(doc, "14.3 Before/After JSON Karşılaştırma", bold=True, size=13)
add_para(
    doc,
    "Her değişiklik için 'Detay' butonuyla değişikliğin öncesi ve sonrası "
    "JSON formatında karşılaştırmalı görüntülenir. Bu, mali denetim açısından "
    "kritik bir özelliktir.",
)
page_break(doc)

# ============================================================
# 15. SSS
# ============================================================
add_heading(doc, "15. Sık Karşılaşılan Sorular", level=1)

add_para(doc, "S: Şifremi unuttum, ne yapmalıyım?", bold=True)
add_para(
    doc,
    "C: Giriş ekranında 'Şifremi unuttum' bağlantısına tıklayın. Yöneticiniz size "
    "yeni geçici şifre tanımlayabilir. İlk girişte sizden zorunlu olarak yeni "
    "şifre belirlemenizi ister.",
)

add_para(doc, "S: Yazıcımdan Türkçe karakterler bozuk çıkıyor?", bold=True)
add_para(
    doc,
    "C: Ayarlar → Yazıcılar sayfasından ilgili yazıcının ayarlarına gidin. "
    "'PC857 Türkçe Encoding' seçeneğinin aktif olduğundan emin olun. Bazı "
    "yazıcılarda ek olarak 'skipInit' bayrağını da etkinleştirmek gerekebilir.",
)

add_para(doc, "S: İnternet bağlantım yok, sistem çalışır mı?", bold=True)
add_para(
    doc,
    "C: Evet. Sistem tamamen çevrimdışı (offline) çalışır. Yalnızca "
    "otomatik güncelleme ve Sentry hata izleme özellikleri internet gerektirir.",
)

add_para(doc, "S: Kapanmış günün siparişine iade yapabilir miyim?", bold=True)
add_para(
    doc,
    "C: Hayır. Z raporu alınıp dönem kapatıldıktan sonra o güne iade kaydı "
    "yapılamaz. Bu, mali güvence açısından bilinçli bir kısıtlamadır.",
)

add_para(doc, "S: Verilerimi başka bilgisayara nasıl taşırım?", bold=True)
add_para(
    doc,
    "C: Manuel yedek alın → ZIP dosyasını yeni bilgisayara taşıyın → "
    "kurulum sonrası 'Dosyadan Geri Yükle' ile yedeği yükleyin. SHA-256 ile "
    "bütünlük doğrulanır.",
)

add_para(doc, "S: Birden fazla yazıcım var, hepsinden çıktı alabilir miyim?", bold=True)
add_para(
    doc,
    "C: Evet. Her kategori için ayrı yazıcı hedefi seçebilirsiniz. Örneğin "
    "yemekler mutfak yazıcısına, içecekler bar yazıcısına otomatik yönlendirilir.",
)
page_break(doc)

# ============================================================
# 16. DESTEK
# ============================================================
add_heading(doc, "16. Destek ve İletişim", level=1)
add_para(doc, "Destek Kanalları", bold=True, size=13)
add_bullet(doc, "📧 E-posta: ilhanavci499@gmail.com")
add_bullet(doc, "🐙 GitHub Issues: github.com/ilhanavc/restoran-pos/issues")
add_bullet(doc, "📦 Sürüm İndirme: github.com/ilhanavc/restoran-pos/releases")

add_para(doc, "Destek Paketi", bold=True, size=13)
add_para(
    doc,
    "Bir sorun yaşadığınızda, Ayarlar → Store Bridge ekranındaki "
    "'Destek Paketi İndir' butonu ile sistem durumu, log ve yapılandırma "
    "bilgilerinin olduğu ZIP dosyasını oluşturup tarafımıza iletebilirsiniz. "
    "Bu dosya sorununun hızlı çözülmesine yardımcı olur.",
)

add_para(doc, "Sürüm Bilgisi", bold=True, size=13)
add_bullet(doc, "Mevcut sürüm: v1.1.0")
add_bullet(doc, "Yayın tarihi: Mayıs 2026")
add_bullet(doc, "Lisans: Akademik Bitirme Projesi")

add_para(doc, "")
add_para(
    doc,
    "Bu kılavuz Restoran POS v1.1.0 sürümü için hazırlanmıştır. "
    "Daha fazla teknik detay için projenin GitHub deposundaki dokümantasyona göz atabilirsiniz.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

# Save
output_path = r"D:\dev\restoran-pos-v3\docs\KULLANIM-KILAVUZU-Restoran-POS.docx"
doc.save(output_path)
print(f"OK — yazıldı: {output_path}")
